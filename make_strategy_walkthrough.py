"""Two-part strategy walkthrough for the team: (1) plain-English story, (2) technical detail + the
'why we chose what we chose' evidence. Output: GSD2T_Strategy_Walkthrough.docx
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

V=json.load(open("v1_flagship.json")); PIT=json.load(open("survivorship_corrected.json"))
CS=json.load(open("concentration_survivorship.json")); CV=json.load(open("concentration_v1.json"))
S=V["summary"]["Fund (full)"]; SPY=V["summary"]["SPY"]; AL=V["alpha"]; NF=V["net_of_fee"]
ISm=V["summary"]["Fund IS (2002-2015)"]; OOS=V["summary"]["Fund OOS (2016-2026)"]
NAVY=RGBColor(0x0B,0x1F,0x3A); GREY=RGBColor(0x55,0x55,0x55); LGREY=RGBColor(0x88,0x88,0x88); GOLD=RGBColor(0xB8,0x86,0x0B); POS=RGBColor(0x2E,0x7D,0x5B)

doc=Document(); doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(11)
def body(t,bold=False,italic=False,size=11,color=None,after=6):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); return p
def bullet(t,lead=None):
    p=doc.add_paragraph(style="List Bullet")
    if lead: r=p.add_run(lead); r.bold=True
    p.add_run(t); return p
def table(headers,rows,style="Light Grid Accent 1"):
    t=doc.add_table(rows=1,cols=len(headers)); t.style=style
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; run=c.paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(10)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=""; run=cs[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(10)
    return t
def pc(x): return f"{x*100:.1f}%"

doc.add_heading("GSD²T — How the Strategy Works",0)
sub=doc.add_paragraph(); r=sub.add_run("A plain-English story, then the technical detail · for the team")
r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GREY
n=doc.add_paragraph(); rn=n.add_run("All performance is SIMULATED (2002–2026), net of trading costs. The fund is fictional.")
rn.italic=True; rn.font.size=Pt(9); rn.font.color.rgb=LGREY

# ===================== PART 1 — THE STORY =====================
doc.add_heading("Part 1 — The strategy in plain English",1)

doc.add_heading("The one-sentence version",2)
body("We buy the stocks that have been going up, we spread our money across about 125 of them so no single "
     "company can hurt us, and — most importantly — we run a system that tells us when to pull money OUT of the "
     "stock market and park it somewhere safe before things get ugly. We make money the way a careful driver wins "
     "a long race: not by being fastest on every straight, but by never crashing.", bold=False)

doc.add_heading("How it works — three simple ideas",2)
body("Idea 1 — Ride the winners.", bold=True, after=2)
body("Stocks that have been rising tend to keep rising for a while. This is called momentum, and it is one of the "
     "most studied patterns in finance. Each month we rank every S&P 500 company by how it did over the past year "
     "and hold the strongest quarter of them (about 125 names), in equal amounts. We are not trying to find the "
     "one genius pick — we hold a broad basket of what is working.")
body("Idea 2 — Read the weather.", bold=True, after=2)
body("This is the real engine. We watch four signals that tell us whether the market environment is calm or stormy: "
     "the 'fear gauge' (VIX), signs of stress in the credit market, the direction of interest rates, and the "
     "overall trend of the market. When the weather is calm, we are fully invested in stocks. When the storm "
     "signals flash, the system automatically cuts our stock exposure — down to as little as 30% — and steps aside. "
     "No emotion, no hesitation; a rule does it for us.")
body("Idea 3 — Don't just sit in cash; hold an umbrella.", bold=True, after=2)
body("When we step out of stocks, we don't hold boring cash. We hold government bonds and gold, which tend to RISE "
     "exactly when stocks are falling. So the money we pull out of the storm actually keeps working for us. This is "
     "the piece we added most recently, and it is what lifts our results above the basic version.")

doc.add_heading("Why this wins: we win by losing less",2)
body("Over the last 24 years the stock market roughly halved twice — once in the 2008 financial crisis and again "
     f"in the 2020 COVID crash (about −46% to −51%). Our strategy's worst fall was about {pc(S['MaxDD'])}. Because we "
     "lose far less in the crashes, we start each recovery from a higher base, and over time that compounds into more "
     f"money with a smoother ride: about {pc(NF['net']['CAGR'])} a year after fees versus the market's {pc(SPY['CAGR'])}, "
     "at roughly two-thirds of the market's bumpiness. We are not trying to beat the market on the good days — we win "
     "by not getting hurt on the bad ones.")

doc.add_heading("A picture to hold in your head",2)
body("Think of a sailor on a long voyage. On calm seas they put up every sail and go fast. When they see a storm "
     "coming, they reef the sails and ride it out — they lose a little speed, but they don't capsize. Our four "
     "weather signals are the sailor's eyes, the exposure dial is the sails, and the bonds-and-gold sleeve is the "
     "ballast that steadies the boat in rough water. The whole point is to still be afloat — and ahead — when the "
     "storm passes.", italic=True)

doc.add_heading("Why we chose what we chose (the short version)",2)
bullet("we hold ~125 stocks instead of a concentrated handful, because our skill is the timing system, not picking "
       "the one winner. A wide basket means no single company blowing up can sink us. We tested holding just 25 "
       "'best ideas' — it looked great in one flattering backtest but fell apart on more honest data, so we don't "
       "trust it.", lead="Broad, not concentrated: ")
bullet("it is the most robust, most researched pattern in markets. We are cleanly implementing a known effect, not "
       "inventing a secret signal we can't defend.", lead="Momentum: ")
bullet("anyone can buy winning stocks; very few have the discipline to step aside in a storm. A computer does it "
       "for us, the same way every time, with no fear or greed. This is our actual edge.", lead="The weather overlay: ")
bullet("we only take a performance cut when we beat the market, and only on new highs. We don't charge hedge-fund "
       "prices for something you can get in and out of any month.", lead="Fair fees: ")

# ===================== PART 2 — TECHNICAL =====================
doc.add_page_break()
doc.add_heading("Part 2 — The technical version",1)

doc.add_heading("Universe & data",2)
body("Universe: the full S&P 500, all 11 sectors (~500 large-cap US stocks). Data: monthly total returns from "
     "yfinance; VIX, IEF/HYG, 10-year yield and the S&P 500 index for the overlay; Fama-French 5 factors + "
     "Momentum and the risk-free rate from the Ken French library; Bloomberg point-in-time membership for the "
     "survivorship test. Back-test window 2002–2026 (2000–2001 reserved as signal warm-up).")

doc.add_heading("Layer 1 — Signal & selection",2)
table(["Item","Detail"],[
 ["Signal","12-minus-1 month price momentum (last 12 months' return, skipping the most recent month)"],
 ["Standardisation","Cross-sectional z-score each month (rank vs the universe)"],
 ["Selection","Top quartile (~125 names)"],
 ["Weighting","Equal weight within the book (~0.9% per name at full exposure)"]])

doc.add_heading("Layer 2 — The macro overlay (the risk dial)",2)
body("Four economically-motivated signals, each z-scored over a rolling 5-year window, averaged into one composite "
     "'risk score', clipped to [−2, +2] and lagged one month. Three are stress signals (turn risk down); one is "
     "trend (turn risk up).", after=4)
table(["Factor","Captures","Direction"],[
 ["VIX","Equity fear / volatility","High → de-risk"],
 ["Credit (IEF − HYG)","Credit-market stress","Widening → de-risk"],
 ["10Y yield change","Rate tightening","Rising → de-risk"],
 ["S&P 500 trend","Market momentum","Up → risk-on"]])
body("The score maps to gross equity exposure:  exposure = clip(0.65 + 0.175 × score, 0.30, 1.00).  "
     "Calm → up to 100% in stocks; stress → as low as 30%.", italic=True, after=4)

doc.add_heading("Layer 3 — Defensive sleeve (the 'V1' refinement)",2)
body("The un-invested portion (1 − exposure) is NOT held in cash — it earns a Treasuries + gold basket (TLT/GLD, "
     "with IEF as the early-period fallback). These assets tend to rise in equity stress, so the risk-off sleeve "
     "contributes rather than just sitting idle. This is the single change that turned the baseline into the V1 "
     "flagship.")

doc.add_heading("Mechanics",2)
bullet("Long-only, no leverage, no shorting.")
bullet("Monthly rebalance; 15 bps round-trip transaction cost applied to all turnover (figures are net of this).")
bullet(f"Average ~{(1-V['ops']['avg_cash'])*100:.0f}% in equities over time, the rest in the defensive sleeve.")

doc.add_heading("Results (simulated, 2002–2026, net of costs, gross of fund fees)",2)
table(["Metric","GSD²T (V1)","S&P 500"],[
 ["CAGR (gross)",pc(S['CAGR']),pc(SPY['CAGR'])],
 ["CAGR (net to investor)",pc(NF['net']['CAGR']),pc(NF['spy']['CAGR'])],
 ["Volatility",pc(S['Vol']),pc(SPY['Vol'])],
 ["Sharpe",f"{S['Sharpe']:.2f}",f"{SPY['Sharpe']:.2f}"],
 ["Max drawdown",pc(S['MaxDD']),pc(SPY['MaxDD'])],
 ["Alpha vs FF5+Momentum",f"+{AL['annualized']*100:.1f}% (t={AL['tstat']:.1f})","—"],
 ["In-sample / out-of-sample Sharpe",f"{ISm['Sharpe']:.2f} / {OOS['Sharpe']:.2f}","—"]])
body(f"Market beta is only {AL['betas']['Mkt-RF']:.2f} — the strategy carries low net market exposure, which is why "
     "it falls so much less in crashes.", italic=True, size=10)

doc.add_heading("Design decisions & the evidence behind them",2)
body("Every choice was tested, not assumed. We judged on out-of-sample results to avoid fooling ourselves.", italic=True, after=4)
bullet("we tested holding 25 / 50 / 125 / 165 names. A concentrated 25-name book looks best in our most optimistic "
       "setup (monthly, current constituents: Sharpe ~1.20) — but it REVERSES on the survivorship-free, point-in-time "
       f"data, where the diversified quartile wins on both Sharpe ({CS['free']['Quartile (~125)']['Sharpe']:.2f} vs "
       f"{CS['free']['Top 25 names']['Sharpe']:.2f}) and drawdown (−21% vs −31%). An edge that only appears in our most "
       "flattering backtest is an artifact, not an edge — so we hold the diversified book.", lead="Diversified (quartile), not concentrated: ")
bullet("we tested inverse-volatility (risk-parity) weighting. It LOWERED the Sharpe "
       f"(to {CV['Quartile + inverse-vol weight']['full']['Sharpe']:.2f} vs {S['Sharpe']:.2f}) because momentum returns "
       "concentrate in higher-volatility winners. Equal weight won on evidence.", lead="Equal weight, not risk-parity: ")
bullet("the top quartile is the standard cut in the academic momentum literature and gives ~125 names — enough to "
       "express the factor cleanly without diluting into the weak half of the distribution.", lead="Top-quartile cut: ")
bullet("a head-to-head bake-off showed the bonds-and-gold sleeve beat a cash sleeve on return, Sharpe AND drawdown, "
       "in and out of sample. We adopted it as the flagship (V1).", lead="Defensive sleeve over cash: ")
bullet("we tested the standard literature enhancements — risk-managed momentum, volatility targeting, low-beta and "
       "quality tilts. NONE improved out-of-sample performance, so we kept the simple design. Discipline, not "
       "complexity.", lead="Enhancements tested and rejected: ")
bullet(f"using current index members overstates results because failed companies drop out. We measured the bias on "
       f"Bloomberg point-in-time data (1,085 names incl. delisted): only ~{abs(PIT['bias']['cagr_drag'])*100:.1f}%/yr. "
       "Small, quantified, disclosed — not assumed.", lead="Survivorship bias corrected: ")
bullet("the same engine gives strong, significant results on the full S&P 500 and even the Dow 30 — evidence of a "
       "real effect, not a single-market fit. We pitch the diversified sector-wide version.", lead="Sector-wide, not tech-only: ")

doc.add_heading("Risk framework",2)
bullet("Primary control: the macro overlay de-grosses to as low as 30% equities in stress — the source of the −21% "
       "drawdown versus the market's −46/−51%.")
bullet("Diversification: ~125 names across 11 sectors, ~1% single-name cap, no leverage, no shorting.")
bullet("Factor discipline: low, controlled exposures; alpha measured against FF5+Momentum with Newey-West standard errors.")
bullet("Defensive convexity: the risk-off sleeve (Treasuries + gold) tends to rise when equities fall.")

doc.add_heading("Honest limitations (what we disclose)",2)
bullet("Returns are net of trading costs but gross of fund fees and taxes; investor-net is shown separately.")
bullet("The defensive sleeve relies on bonds/gold being defensive; 2022 broke that briefly (bonds fell with stocks), "
       "though the out-of-sample period including 2022 still beats the market. Part of the historical benefit came "
       "from a multi-decade bond tailwind that will not fully repeat.")
bullet("All performance is simulated; no live track record exists.")

d=doc.add_paragraph(); rd=d.add_run("Disclaimer: Fictional pitch for the ESADE Asset Management course. Not investment advice. "
    "All performance is simulated; past performance does not indicate future results.")
rd.italic=True; rd.font.size=Pt(8); rd.font.color.rgb=LGREY

doc.save("GSD2T_Strategy_Walkthrough.docx")
print("Saved GSD2T_Strategy_Walkthrough.docx")
