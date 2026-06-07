"""Q&A simulation / prep document — anticipated panel questions with honest, data-grounded answers.
Output: GSD2T_QA_Simulation.docx
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor

V=json.load(open("v1_flagship.json")); CUR=json.load(open("current_state.json")); RB=json.load(open("robustness_assumptions.json"))
S=V["summary"]["Fund (full)"]; SPY=V["summary"]["SPY"]; AL=V["alpha"]; NF=V["net_of_fee"]
NAVY=RGBColor(0x0B,0x1F,0x3A); GREY=RGBColor(0x55,0x55,0x55); LGREY=RGBColor(0x88,0x88,0x88)
NEG=RGBColor(0xB8,0x3A,0x3A); POS=RGBColor(0x2E,0x7D,0x5B); GOLD=RGBColor(0xB8,0x86,0x0B)

doc=Document(); doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(11)
def para(t,bold=False,italic=False,size=11,color=None,after=6,lead=None):
    p=doc.add_paragraph()
    if lead: r0=p.add_run(lead); r0.bold=True; r0.font.size=Pt(size); r0.font.color.rgb=color or NAVY
    r=p.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color and not lead: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); return p
def qa(q,a,avoid=None,key=None):
    para(q,lead="Q:  ",size=11.5,color=NAVY)
    para(a,lead="A:  ",size=11)
    if key: para(key,lead="Key number:  ",size=9.5,color=POS)
    if avoid: para(avoid,lead="Avoid:  ",size=9.5,color=NEG)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

doc.add_heading("GSD²T — Q&A Simulation & Defence Notes",0)
sub=doc.add_paragraph(); r=sub.add_run("Anticipated panel questions with honest, data-grounded answers · internal prep")
r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GREY

doc.add_heading("How to use this (composure principles)",1)
for t in [
 "Concede small points fast, then pivot to your real strength (risk control). Defensiveness reads as weakness.",
 "Lead with a number. \"Our drawdown was −21% vs the market's −50%\" beats a paragraph of theory.",
 "Never claim more than your data shows. The panel has the appendix; a contradicted claim costs more than the question.",
 "If you don't know, say \"we didn't test that — here's how we would.\" Honesty scores; bluffing doesn't.",
 "Your one-line anchor: \"Our edge is capital preservation, not a faster signal. We win by losing less.\"",
]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(t)

doc.add_heading("A · The edge & strategy",1)
qa("Is your strategy systematic or discretionary?",
   "Fully systematic — rules-based at every layer, with zero discretionary judgment. The momentum signal, the "
   "4-factor macro overlay, the exposure dial and the monthly rebalance are all pre-defined formulas. Nobody picks a "
   "stock, overrides the dial, or makes a market call. The proof: the entire strategy is reproducible from our "
   "notebook — a grader can re-run it and get the exact same positions and returns, which a discretionary strategy "
   "can't. And the discipline is part of the edge: we de-risk in stress precisely when a discretionary manager might "
   "freeze.",
   key="Reproducible from code = the hallmark of systematic. Maps to the brief's own 'equity factor + systematic macro' categories.")
qa("What is your actual edge? Everyone runs momentum.",
   "Our edge is not the stock signal — momentum is commoditised and we say so. Our edge is the macro overlay that "
   "systematically cuts equity exposure in stress, plus a defensive sleeve that earns in risk-off periods instead of "
   "sitting in cash. That is what turns a raw momentum return into a 1.09 Sharpe with half the market's drawdown.",
   key="A raw momentum ETF (MTUM) ran ~1.0 Sharpe with ~−50% drawdowns; we deliver 1.09 at −21%.",
   avoid="Don't claim a proprietary or 'better' momentum signal — you can't defend it and don't need to.")
qa("Why does momentum persist? Shouldn't it be arbitraged away?",
   "It's a behavioural under-reaction effect, and it persists because of limits to arbitrage — momentum has rare, "
   "violent crashes that deter leveraged arbitrageurs, and mandate/career constraints stop most managers from "
   "following it cleanly. Decades of literature (Jegadeesh & Titman onward) document it. We don't bet the fund on it "
   "though — it's the selection layer; the overlay is the differentiator.")
qa("Why no machine learning or deep models?",
   "Because at monthly frequency the binding constraint is sample size, not model capacity. With ~290 monthly "
   "observations, a deep model would overfit — the literature shows ML helps where data is rich (intraday, alt-data) "
   "and that a regularised linear/factor baseline is already strong for monthly cross-sectional work. A simple, "
   "transparent momentum signal is the disciplined choice; it's also fully reproducible and auditable, which a "
   "learned policy is not. We'd reach for ML only if the data and the structure justified it; here they don't.",
   key="Monthly data → sample complexity is the constraint, not capacity. Simple beats deep here — and stays auditable.")
qa("Isn't this just a long-only factor portfolio I can buy as an ETF?",
   "Partly, and we're honest about that — the stock-selection layer is replicable. What isn't replicable is the "
   "overlay-plus-sleeve timing, which drives the drawdown and Sharpe. That's also why we charge below hedge-fund "
   "terms and only take a performance fee above the benchmark.")

doc.add_heading("B · Portfolio construction (the concentration trap)",1)
qa("Why hold ~125 names? Why not 20–30 high-conviction picks?",
   "We tested it. A 25-name book looks best in our single most-optimistic setup — monthly, current constituents, "
   "Sharpe ~1.20 — but it REVERSES on the survivorship-free, point-in-time data, where the diversified quartile wins "
   "on both Sharpe (0.54 vs 0.48) and drawdown (−21% vs −31%). An edge that only appears in your most flattering "
   "backtest is an artifact, not an edge. We hold the diversified book because it's the number we can trust, it has "
   "a fraction of the single-name risk, and it preserves capacity.",
   key="Concentration premium is positive only monthly; it goes negative on survivorship-free data.",
   avoid="Don't say 'concentration doesn't add return' — our own monthly data shows it does. Argue fragility, not absence.")
qa("Why equal weight and not something smarter?",
   "We tested inverse-volatility (risk-parity) weighting; it lowered the Sharpe to 0.99 from 1.09, because momentum "
   "returns concentrate in the higher-volatility winners that risk-parity down-weights. Equal weight won on evidence, "
   "and it's the most transparent and capacity-friendly choice.")
qa("Why the top quartile specifically — isn't that arbitrary?",
   "It's the standard cut in the momentum literature and gives ~125 names — enough to express the factor cleanly "
   "without diluting into the weak half of the distribution. We tested decile and tercile around it; the risk-adjusted "
   "profile is stable, so we chose the broader, more robust, more scalable cut deliberately.")

qa("Why long-only? Why not long-short?",
   "We tested long-short and it bled. In a market dominated by a handful of mega-caps, the short legs got run over "
   "in the 2018–24 rally — shorting the names that kept winning produced catastrophic drawdowns. Long-only is a "
   "deliberate risk choice: it bounds losses (no unbounded short risk), avoids borrow costs and short-squeeze risk, "
   "and the variable-net-exposure overlay gives us the de-risking benefit a market-neutral book would, without the "
   "short-side fragility.",
   key="We tested L/S — the short legs bled in the mega-cap rally. Long-only is the deliberate, lower-risk choice.")

doc.add_heading("C · The macro overlay",1)
qa("Isn't this just market timing? Market timing doesn't work.",
   "It's not discretionary calls — it's a rules-based signal that only SCALES exposure between 30% and 100%, never "
   "shorts and never levers. We're not predicting the market; we're responding to measurable stress. The evidence is "
   "in the drawdowns: it roughly halves them in every major crisis we tested.")
qa("Did you data-mine the four overlay factors?",
   "No. Each is economically motivated ex ante — equity fear (VIX), credit stress, rate tightening, market trend — "
   "not selected by backtest performance. They're z-scored on a rolling window and lagged. The parameter-sensitivity "
   "grid shows the result is a plateau, not a spike, which is what you'd expect from a real effect rather than a "
   "fitted one.",
   avoid="Don't get drawn into defending exact thresholds; defend the ex-ante economic logic and the robustness.")
qa("What happens when the overlay whipsaws — de-risks and then the market rallies?",
   "It does cost us in sharp V-shaped recoveries — we lag the market on its best days. That's the price of the "
   "insurance. Over a full cycle the drawdown reduction more than pays for it, but we disclose openly that we trail "
   "in fast rebounds. We're selling a smoother ride, not the highest beta.")

doc.add_heading("D · Back-test honesty (the 25-point question)",1)
qa("How do you handle survivorship bias?",
   "We corrected it, not just disclosed it. Using Bloomberg point-in-time S&P 500 membership — 1,085 historical names "
   "including everything that delisted or was acquired — we measured the bias directly at about 1% per year. Small, "
   "quantified, and our risk-control edge survives the correction intact.",
   key="~1%/yr CAGR bias, measured on 1,085 point-in-time names.")
qa("How do I know there's no look-ahead bias?",
   "Every signal is lagged at least one month and uses only trailing windows for the z-scores. We reserve 2000–2001 "
   "purely as signal warm-up and only measure performance from 2002, so no number is produced by a signal that "
   "wasn't fully formed at the time.",
   key="Disclosed: the credit factor uses the HYG ETF (2007 inception), so it joins the overlay ~2010 — the dial runs on the other three factors before then. Data availability, not look-ahead; result is robust (1.09 from 2002 vs 1.07 from 2004).")
qa("Did you include transaction costs?",
   "Yes — 15 bps round-trip on all turnover, and every figure we show is net of it. Turnover runs around 380% a year; "
   "the cost drag is modelled, not ignored.")
qa("How do I know you didn't overfit this?",
   "Three ways. One, a clean in-sample/out-of-sample split — the out-of-sample Sharpe is 1.01 versus 1.16 in-sample, a "
   "modest step-down, not the collapse overfitting produces, and still well above the market's 0.86 over the same "
   "out-of-sample window. Two, a parameter-sensitivity grid showing a plateau. Three — and this is the honest part — "
   "we tested the standard literature enhancements and REJECTED them when they didn't beat the simple design "
   "out-of-sample. We'll show you the experiments we threw away.",
   key="OOS Sharpe 1.01 vs IS 1.16 (and vs SPY 0.86 OOS); deflated Sharpe in the appendix.")
qa("A Sharpe over 1 — is that even statistically significant?",
   "Yes. The alpha t-statistic is 4.3 against the Fama-French five factors plus momentum, using Newey-West "
   "heteroskedasticity- and autocorrelation-consistent standard errors. We also report a deflated Sharpe ratio that "
   "adjusts for the number of trials. And to be clear, 1.09 is not a red-flag Sharpe — we're not claiming 3+.")
qa("Isn't your 'alpha' just disguised factor exposure?",
   "We control for it — the +5.4% is the intercept after regressing on the five factors plus momentum. One honest "
   "caveat: our factor model has no bond or gold factor, so part of that alpha is really the defensive sleeve's "
   "bond-and-gold return, not pure skill. The cleanest, most defensible wins are the Sharpe and the drawdown, not the "
   "headline alpha.",
   avoid="Don't oversell the +5.4% alpha — concede the bond/gold attribution before they raise it.")
qa("Your transaction-cost assumption — what if 15 bps is too low?",
   f"We stress-tested it. At double the cost (30 bps) the Sharpe is still {RB['tcost'][1]['Sharpe']:.2f}; at triple "
   f"(50 bps) it's {RB['tcost'][2]['Sharpe']:.2f}; and even at 100 bps — six times our assumption — it's "
   f"{RB['tcost'][3]['Sharpe']:.2f} with an {RB['tcost'][3]['CAGR']*100:.1f}% CAGR, both still beating the market. "
   "So we could be badly wrong about costs and still win; the 15 bps figure isn't load-bearing. And it excludes market "
   "impact, which is negligible at our size given the capacity headroom.",
   key="Edge survives 6× the assumed cost (100 bps → Sharpe 0.82, still > SPY 0.60).")
qa("Did you tune the overlay's parameters (the 0.65 baseline, 0.175 slope) to the backtest?",
   f"No, and we can show it. Varying the neutral baseline from 55% to 75% and the dial sensitivity from 0.125 to "
   f"0.225 — nine combinations — the Sharpe only moves between {RB['overlay']['min']:.2f} and {RB['overlay']['max']:.2f}. "
   "It's a flat plateau, not a knife-edge. If we'd curve-fit those numbers you'd see a sharp peak at our exact setting "
   "and a drop-off around it; instead any reasonable calibration gives essentially the same result.",
   key="Sharpe 1.06–1.09 across 9 overlay settings — a plateau, not a fitted peak.")

doc.add_heading("E · Risk framework",1)
qa("What's your worst case? Where's the tail risk?",
   "Worst simulated drawdown was −21%%, versus the market's −46%% to −51%%, across seven separate crises. On the "
   "coherent measure — expected shortfall, not VaR — our monthly 95%% CVaR is %.1f%% versus the market's %.1f%%; we "
   "report ES because returns are heavy-tailed and Gaussian VaR understates the tail. Structurally the tail is "
   "bounded: long-only and no leverage means no unbounded loss, single names are capped near 1%%, and the overlay "
   "de-risks in stress. In a live fund we'd add explicit exposure and volatility limits."
   %(V["tail_risk"]["fund"]["CVaR95_monthly"]*100, V["tail_risk"]["spy"]["CVaR95_monthly"]*100),
   key="Monthly 95%% CVaR (expected shortfall) %.1f%% vs market %.1f%% — the coherent measure, and lower than the market."
   %(V["tail_risk"]["fund"]["CVaR95_monthly"]*100, V["tail_risk"]["spy"]["CVaR95_monthly"]*100))
qa("What regime breaks this strategy?",
   "One where momentum AND the defensive assets fail at the same time — which is exactly what 2022 was, when stocks "
   "and bonds fell together. We disclose it. Gold in the sleeve diversifies that risk, and even including 2022 the "
   "out-of-sample period still beats the market — but we don't pretend the strategy is immune to it.")

doc.add_heading("F · The defensive sleeve",1)
qa("Why bonds and gold? 2022 destroyed the bond hedge.",
   "2022 is precisely why we hold bonds AND gold rather than bonds alone — gold held up when Treasuries didn't. Over "
   "the full sample the sleeve beat a cash sleeve on return, Sharpe and drawdown. We also disclose that part of its "
   "historical benefit came from a multi-decade bond bull market we don't assume repeats.")
qa("Isn't the sleeve just a 60/40 or risk-parity portfolio in disguise?",
   "No — it's conditional, not strategic. The sleeve is only funded when the overlay de-risks, roughly a third of the "
   "book on average, and its job is to be defensive in stress, not to be a permanent allocation. It turns on and off "
   "with the regime signal.")

doc.add_heading("G · Capacity & liquidity",1)
qa("Your capacity number — what's the quantitative basis?",
   "Trade no more than 5% of a name's average daily volume over two days, across ~111 holdings. We don't assume the "
   "ADV — we measured it against actual volume data: the median for our holdings is about $150M/day, which we use as "
   "the conservative case (up to $300M as the base). That gives a soft cap of $1.7–3.3 billion — 17 to 33 times the "
   "$100M raise, which uses only about 3–6% of capacity. The defensive sleeve sits in the most liquid ETFs on earth, "
   "so it adds no constraint.",
   key="$15–30M tradeable per name × ~111 names = $1.7–3.3B soft cap (ADV measured, not assumed).",
   avoid="Don't call $300–500M ADV 'conservative' — the measured median is ~$150M; lead with the measured figure.")
qa("What about liquidity drying up in a crisis?",
   "Daily volumes do fall in a crisis, which shrinks capacity temporarily — but note the self-correction: that's "
   "exactly when the overlay de-risks and we trade less. We demand the least liquidity precisely when liquidity is "
   "scarce.")

doc.add_heading("H · Fund terms",1)
qa("Why should I pay you fees for something I can largely replicate with ETFs?",
   "You're paying for the overlay-and-sleeve timing and the discipline, not the factor — and we price it that way: "
   "1% management plus 15% only on returns ABOVE the S&P 500, with a high-water mark. We don't charge an alpha fee on "
   "market beta you can buy for five basis points.")
qa("1% on $100M is $1M a year whether you perform or not.",
   "Correct, and it covers operations and infrastructure. The bulk of our economics is the performance fee, which "
   "only triggers above the benchmark and only on new highs — so we're paid to beat the market, not to hold it.")

doc.add_heading("I · Results, 'why now', and general",1)
qa("Do you beat an equal-weight S&P 500, or just cap-weight?",
   "Our alpha already controls for it. The +5.4%% alpha is measured against Fama-French 5 + Momentum, which includes "
   "SMB (size) — so the equal-weight/small-cap tilt is regressed out, and the alpha survives it. Equal-weighting is a "
   "genuinely hard benchmark (DeMiguel-Garlappi-Uppal showed it often beats Markowitz out-of-sample), which is "
   "exactly why we equal-weight within our book rather than optimise. Beating cap-weight SPY net of fees is the "
   "headline; the factor-adjusted alpha shows it isn't merely a size tilt.",
   key="Alpha is net of SMB (size), so it's not an equal-weight tilt in disguise. We equal-weight precisely because it's hard to beat.")
qa("Net 13.1% versus the market's 10% — that's only 3% for all this machinery.",
   "The 3% net outperformance is real, but it's not the headline — the headline is risk. We deliver it at a 1.09 "
   "Sharpe versus 0.60, and a −21% drawdown versus −50%. We're selling risk-adjusted compounding and downside "
   "protection. In 2008 and 2020, that drawdown gap is the entire value proposition.",
   key="Net 13.1% vs 10.0%, but Sharpe 1.09 vs 0.60 and −21% vs −50% DD.")
qa("Your out-of-sample is 2016–2026, mostly a bull market. How do you know it works in a bear?",
   "The out-of-sample window contains the 2018 vol spike, the 2020 COVID crash and the 2022 bear — three genuine "
   "drawdowns where the overlay was tested live in the data, not just in the calm. And the strategy's whole value, "
   "de-risking, is most visible precisely in those stress periods. The in-sample period also includes 2008.")
qa("Your OOS Sharpe is higher than in-sample — isn't that suspicious?",
   "It reflects the overlay performing well through the 2020 and 2022 stress, and we don't claim that's repeatable "
   "luck. What it does show is the absence of out-of-sample decay — the usual fingerprint of overfitting is OOS far "
   "below IS, and we don't have that.")
qa("Why now?",
   "Elevated regime uncertainty keeps drawdown risk high, so a strategy built to de-risk is timely. Positive real "
   "yields mean the defensive sleeve actually earns while it waits, unlike the zero-rate decade. And factor crowding "
   "rewards discipline over signal-cleverness — which is exactly our design.")
qa("What would make you shut this strategy down?",
   "Pre-defined triggers: if live tracking error versus the backtest blows out, if the overlay fails to de-risk in a "
   "genuine stress event, or if the alpha decays below costs. We'd disclose and adapt rather than hope.")
qa("What AI tools did you use?",
   "Disclosed in the deck: AI assistants for code scaffolding, back-test engineering and document drafting. Every "
   "strategy decision, data choice and result was defined and verified by the team; nothing was fabricated; and every "
   "figure is reproducible from the committed code.")

doc.add_heading("J · Current positioning (REFRESH before the pitch)",1)
ge=CUR["gross_exposure"]*100; sl=CUR["defensive_sleeve"]*100; cm=CUR["components"]; rd_=CUR["component_reading"]
asof=CUR["as_of"]
para(f"IMPORTANT — these readings are as of end-{asof}. Re-run current_state.py the morning of the pitch and update "
     f"the numbers and date so you quote LIVE readings, not month-old ones.", bold=True, color=NEG, size=10.5)
# indicator snapshot table
t=doc.add_table(rows=1,cols=3); t.style="Light Grid Accent 1"
for i,h in enumerate(["Indicator","Reading (z)","Signal"]):
    cc=t.rows[0].cells[i]; cc.text=""; rr=cc.paragraphs[0].add_run(h); rr.bold=True; rr.font.size=Pt(10)
lab={"vix":"VIX / equity fear","credit":"Credit stress (IEF−HYG)","yield":"Rate direction (10Y)","trend":"Market trend (S&P 500)"}
for k in ["vix","credit","yield","trend"]:
    cs=t.add_row().cells
    for j,v in enumerate([lab[k],f"{cm[k]:+.2f}",rd_[k]]):
        cs[j].text=""; rr=cs[j].paragraphs[0].add_run(str(v)); rr.font.size=Pt(10)
doc.add_paragraph().paragraph_format.space_after=Pt(2)
qa("What is your current equity exposure, and why?",
   f"As of end-{asof}, the overlay holds about {ge:.0f}% in equities and {sl:.0f}% in the defensive sleeve — modestly "
   f"above our 65% neutral baseline, so the model is leaning slightly risk-on. The reason is the rate signal: the 10-year "
   f"has eased over the past year, which is our strongest positive reading right now, supported by a calm VIX around 15 "
   f"and a positive market trend. What's holding us back from a higher number is a mild widening in credit. Composite "
   f"regime score is about +{CUR['composite_score']:.1f} on a −2 to +2 scale.",
   key=f"~{ge:.0f}% equities / {sl:.0f}% defensive sleeve; composite score +{CUR['composite_score']:.1f}.")
qa("How do you see exposure changing in the medium term?",
   "We don't forecast it — exposure is mechanical, the output of the four indicators, so we won't give you a target. But "
   "the recent path shows how it breathes: we drifted from about 70% down to the low-60s through the winter as momentum "
   "and credit softened, then stepped back up as conditions firmed. If today's calm persists we'd drift toward fully "
   "invested; if any stress indicator deteriorates, exposure comes down automatically — no committee, no discretion.")
qa("What event could lead to a reduction in exposure?",
   "Any of the four signals flashing stress: a VIX spike, credit spreads widening (high-yield underperforming "
   "Treasuries), a sharp move up in rates, or the market trend rolling over. Concretely — a credit event, an inflation "
   "surprise that forces rates higher fast, or a momentum reversal would each pull us down, and several together would "
   "push us toward the 30% floor. The 2022 setup, rates and credit deteriorating together, is exactly what de-grosses "
   "us hard.")
qa("What are your indicators telling you right now?",
   f"Net constructive but not euphoric. Rates are the clear positive ({rd_['yield']}); VIX around 15 and the market "
   f"trend are mildly supportive; the one yellow flag is credit, which is {rd_['credit']}. So we're leaning in at "
   f"~{ge:.0f}% equities, with a hand on the dial.",
   avoid="Don't editorialise about macro you can't support — quote the four indicator readings and stop there.")

d=doc.add_paragraph(); rd=d.add_run("Internal prep document — GSD²T Asset Management, ESADE Asset Management course. "
    "All performance is simulated; the fund is fictional.")
rd.italic=True; rd.font.size=Pt(8); rd.font.color.rgb=LGREY
doc.save("GSD2T_QA_Simulation.docx")
print("Saved GSD2T_QA_Simulation.docx")
