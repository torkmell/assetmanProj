"""Step-by-step Bloomberg extraction guide for teammates (survivorship-corrected S&P 500 data).
Covers: yearly membership snapshots, monthly membership snapshots, and monthly prices.
Output: Bloomberg_Data_Extraction_Guide.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F, 0x2D, 0x50)
GREY = RGBColor(0x66, 0x66, 0x66)
LGREY = RGBColor(0x88, 0x88, 0x88)
GOLD = RGBColor(0xB8, 0x86, 0x0B)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

def body(t, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(t, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(t)
    else:
        p.add_run(t)
    return p

def numbered(t, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(t)
    else:
        p.add_run(t)
    return p

def code(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = "Consolas"; r.font.size = Pt(9.5)
    r.font.color.rgb = NAVY
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""
            run = cs[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(10)
    return t

# ---------------- TITLE ----------------
doc.add_heading("Bloomberg Data Extraction — Step-by-Step Guide", level=0)
sub = doc.add_paragraph()
r = sub.add_run("GSD2T Asset Management · ESADE · Survivorship-corrected S&P 500 universe")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
body("Goal: rebuild the S&P 500 as it actually looked at each point in the past (including companies "
     "that have since left the index), so our backtest is not biased by only using today's winners. "
     "This needs TWO things from the Bloomberg Terminal: (1) the index membership over time, and "
     "(2) monthly total-return prices for every name. This guide shows exactly how to pull both.",
     space_after=10)

# ---------------- THE BIG PICTURE ----------------
doc.add_heading("0. The big picture (read this first)", 1)
body("There are two separate pulls, at two different frequencies. Do not confuse them:", bold=True)
table(["Pull", "Frequency", "What it gives", "Bloomberg tool"], [
    ["A. Membership", "Yearly OR Monthly", "Which stocks were in the S&P 500 on each date", "BDS or BQL"],
    ["B. Prices", "Monthly (always)", "Each stock's monthly total return (~316 points/stock)", "BDH"],
])
body("Important: only the MEMBERSHIP is sampled yearly/monthly. The PRICES are always monthly "
     "(~316 observations per stock, 2000–2026) — exactly like a normal backtest. We never use yearly prices.",
     italic=True, color=GREY)
body("Recommended order: do Pull A (membership) first and send it to the analyst. They will return a "
     "short list of the 'leaver' tickers we don't already have, so Pull B (prices) only needs that short "
     "list instead of the whole universe — saving terminal time.", space_after=10)

# ---------------- SETUP ----------------
doc.add_heading("1. Set-up (once)", 1)
numbered("On the Bloomberg PC, open Excel. Check there is a \"Bloomberg\" tab in the ribbon.")
numbered("Make sure you are logged into the Terminal (the Excel add-in pulls data through your live session).")
numbered("If the Bloomberg tab is missing: on the Terminal type  WAPI <GO>  and reinstall the Excel Add-in, then reopen Excel.")
body("All formulas below are typed into Excel cells, not into the black Terminal screen. The Terminal "
     "screen is only used for look-ups and to launch the wizards.", italic=True, color=GREY)

# ---------------- PULL A ----------------
doc.add_heading("2. Pull A — Index membership", 1)
body("We need to know which stocks were in the S&P 500 on each historical date. Pick ONE of the two "
     "methods below. Yearly is simplest; monthly is more precise (drops a leaver in the right month "
     "instead of up to a year late). Either is acceptable — monthly is preferred if BQL works on your terminal.",
     space_after=8)

# ---- 2A YEARLY ----
doc.add_heading("2A. YEARLY membership snapshots (simplest — uses BDS)", 2)
body("This produces one snapshot per year (about 26 in total).", color=GREY, size=10)
numbered("In a worksheet, type year-end dates down column A, one per row:")
code("A1:  20001231\nA2:  20011231\nA3:  20021231\n ...  (continue every year)\nA26: 20251231")
numbered("In cell B1, reference the date in A1:")
code('=BDS("SPX Index","INDX_MWEIGHT_HIST","END_DATE_OVERRIDE="&A1)')
numbered("Press Enter. It \"spills\" a list of tickers + index weights as of 31 Dec 2000 below B1.")
numbered("Each snapshot is ~500 rows, so they will overwrite each other if stacked. Give each year its "
         "OWN block: put the next formula far to the right (e.g. D1 referencing A2, G1 referencing A3, and so on), "
         "OR put each year on its own worksheet tab. Do not stack them in one column.")
numbered("Repeat for all years 2000–2025.")
body("If INDX_MWEIGHT_HIST returns blank, try the field  INDX_MEMBERS  with the same END_DATE_OVERRIDE.",
     italic=True, color=GREY, size=10)

# ---- 2B MONTHLY ----
doc.add_heading("2B. MONTHLY membership snapshots (more precise — uses BQL, one formula)", 2)
body("A monthly snapshot would be ~316 separate BDS blocks (painful). Instead, BQL (Bloomberg Query "
     "Language) returns the WHOLE monthly history in a SINGLE formula. This is the preferred method.",
     color=GREY, size=10)
numbered("On the Terminal, type  BQLX <GO>  to open the BQL Builder (it writes the syntax for you).")
numbered("Universe: type  SPX Index  and choose  members.")
numbered("Set the members to return over a DATE RANGE, frequency = Monthly, from 2000-01-31 to 2026-05-31.")
numbered("The Builder generates an Excel formula similar to this (verify the exact text in the Builder):")
code('=BQL("members(\'SPX Index\',dates=RANGE(2000-01-31,2026-05-31),frq=M)","id().name")')
numbered("Paste it into an empty Excel cell. It returns one long table of  date + ticker  — every member, "
         "every month — in one shot. No need to manage 316 blocks.")
body("If BQL is not available or misbehaves on your terminal, fall back to QUARTERLY BDS instead: same as "
     "method 2A but use quarter-end dates (20000331, 20000630, 20000930, 20001231, ...). That is 107 blocks "
     "and the analyst can fill the in-between months in code — index membership only changes on those "
     "rebalance dates, so quarterly is effectively exact.", italic=True, color=GREY, size=10)

# ---------------- PULL B ----------------
doc.add_heading("3. Pull B — Monthly total-return prices (BDH)", 1)
body("For every ticker that appears in the membership list (including the leavers), pull monthly "
     "total-return prices. Total return = dividend-adjusted, which matches our methodology.", space_after=8)
body("Single-stock formula (the template):", bold=True, space_after=2)
code('=BDH("AAPL US Equity","TOT_RETURN_INDEX_GROSS_DVDS","1/1/2000","5/31/2026","Per=M","Dts=S")')
table(["Setting", "Meaning"], [
    ['"Per=M"', "Monthly frequency (this is what makes it monthly — about 316 rows per stock)"],
    ['"Dts=S"', "Show the date column"],
    ["TOT_RETURN_INDEX_GROSS_DVDS", "Dividend-adjusted total-return index (use PX_LAST only if this is blank)"],
])
body("Do the whole list at once with the wizard (recommended — no typing 500 formulas):", bold=True, space_after=2)
numbered("Bloomberg ribbon → Spreadsheet Builder → Historical Data (BDH).")
numbered("Securities: paste the entire ticker list (add \" US Equity\" to each, e.g.  AAPL US Equity).")
numbered("Field: TOT_RETURN_INDEX_GROSS_DVDS.")
numbered("Periodicity: Monthly.   Date range: 1/1/2000 to 5/31/2026.")
numbered("Finish — it lays out the full grid: dates down the rows, tickers across the columns, one value per cell.")

# ---------------- DELISTED ----------------
doc.add_heading("4. Handling \"leaver\" / delisted tickers", 1)
bullet("When a company left the index (acquired or bankrupt), Bloomberg keeps it but may need its old "
       "ticker, e.g.  LEH US Equity  (Lehman). ")
bullet("If a name returns blank, type the company name into the Terminal search bar to find its exact "
       "historical ticker, then re-pull that one.")
bullet("Most US large-caps resolve fine; usually only a handful need a manual look-up.")

# ---------------- FREEZE & SAVE ----------------
doc.add_heading("5. Freeze and save (do NOT skip this)", 1)
body("BDS / BDH / BQL cells are LIVE — they turn into #N/A the moment the file is opened off the "
     "Bloomberg PC. You must freeze them to plain values before sending.", bold=True, color=GOLD)
numbered("Select all the data → Copy → Paste Special → Values. (This replaces formulas with the numbers.)")
numbered("File → Save As → CSV UTF-8 (.csv).")
numbered("Do this for both the membership file and the prices file.")

# ---------------- DELIVERABLES ----------------
doc.add_heading("6. What to send back (file formats)", 1)
table(["File", "Layout", "Columns / shape"], [
    ["membership.csv", "One row per member per snapshot", "date, ticker"],
    ["prices.csv", "Matrix", "dates down the rows, tickers across the top, one value per cell"],
])
body("Send the membership file FIRST. The analyst will diff it against the prices we already have and "
     "reply with the exact short list of leaver tickers — then you only pull prices for those.",
     italic=True, color=GREY)

# ---------------- QUICK CHECKLIST ----------------
doc.add_heading("7. Quick checklist", 1)
for t in [
    "Excel open on the Bloomberg PC, logged in, Bloomberg ribbon visible.",
    "Pull A: membership — yearly (2A, BDS) or monthly (2B, BQL). One row per member per date.",
    "Send membership.csv to the analyst; receive the leaver-ticker list back.",
    "Pull B: monthly prices (BDH wizard) for the required tickers, 2000-01 to 2026-05.",
    "Freeze everything to Values (Paste Special), save as CSV UTF-8.",
    "Send membership.csv + prices.csv.",
]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(t)

# ---------------- FOOTER ----------------
d = doc.add_paragraph()
rd = d.add_run("Internal working document — GSD2T Asset Management, ESADE Asset Management course. "
               "Data for a fictional pitch / academic backtest only.")
rd.italic = True; rd.font.size = Pt(8); rd.font.color.rgb = LGREY

doc.save("Bloomberg_Data_Extraction_Guide.docx")
print("Saved Bloomberg_Data_Extraction_Guide.docx")
