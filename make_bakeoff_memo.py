"""Word memo summarising the strategy-refinement bake-offs (V1 defensive sleeve win + improvements test).
Reads bakeoff_variants.json and bakeoff_improvements.json. Output: Strategy_Refinement_Findings.docx
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches

V = json.load(open("bakeoff_variants.json"))
W = json.load(open("bakeoff_improvements.json"))
GREY=RGBColor(0x55,0x55,0x55); LGREY=RGBColor(0x88,0x88,0x88); GOLD=RGBColor(0xB8,0x86,0x0B); POS=RGBColor(0x2E,0x7D,0x5B)

doc=Document(); doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(11)
def body(t,bold=False,italic=False,size=11,color=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return p
def bullet(t,lead=None):
    p=doc.add_paragraph(style="List Bullet")
    if lead: r=p.add_run(lead); r.bold=True
    p.add_run(t); return p
def af(o):
    return "—" if o.get("alpha") is None else f"{o['alpha']*100:+.1f}% ({o['alpha_t']:.1f})"
def perf_table(data, ref_key=None, hi_key=None):
    hdr=["Variant","CAGR","Vol","Sharpe","Max DD","Alpha (t)","IS Shp","OOS Shp","Cash"]
    t=doc.add_table(rows=1,cols=len(hdr)); t.style="Light Grid Accent 1"
    for i,h in enumerate(hdr):
        c=t.rows[0].cells[i]; c.text=""; run=c.paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(8.5)
    for name,o in data.items():
        f=o["full"]; cells=t.add_row().cells
        vals=[name,f"{f['CAGR']*100:.1f}%",f"{f['Vol']*100:.1f}%",f"{f['Sharpe']:.2f}",f"{f['MaxDD']*100:.0f}%",
              af(o),f"{o['is']['Sharpe']:.2f}",f"{o['oos']['Sharpe']:.2f}",f"{o['avg_cash']*100:.0f}%"]
        for i,v in enumerate(vals):
            cells[i].text=""; run=cells[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(8.5)
            if (hi_key and name==hi_key) or (i==0 and name==ref_key): run.bold=True
    return t

doc.add_heading("Strategy Refinement — Bake-off Findings & Recommendation",0)
sub=doc.add_paragraph(); r=sub.add_run("GSD2T Asset Management · ESADE Asset Management · internal research memo")
r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GREY
n=doc.add_paragraph(); rn=n.add_run("All performance SIMULATED, monthly, 2002–2026, net of 15 bps trading costs, GROSS of fund fees. "
    "Sharpe is excess-of-risk-free. IS = in-sample 2002–2015; OOS = out-of-sample 2016–2026.")
rn.italic=True; rn.font.size=Pt(9); rn.font.color.rgb=LGREY

doc.add_heading("1. Objective",1)
body("The team asked two questions: (a) can we put the de-risked 'cash' sleeve to work, and (b) are there "
     "weighting or signal changes that improve risk-adjusted performance while keeping a defensible, "
     "literature-backed narrative. We ran two controlled bake-offs on the live monthly engine, holding "
     "everything else fixed and changing one design choice at a time, and judged each on out-of-sample "
     "Sharpe (not in-sample, to avoid rewarding overfitting).")

doc.add_heading("2. Test 1 — Variant bake-off (how we found V1)",1)
body("Starting from the current strategy (V0: equal-weight top-quartile momentum, macro overlay, cash sleeve), "
     "we tested four ideas: a productive defensive sleeve, inverse-volatility weighting, an asymmetric overlay "
     "(less cash in calm regimes), and different concentration levels.")
perf_table(V, ref_key="V0 Baseline (current)", hi_key="V1 Defensive sleeve")
body("Findings:", bold=True)
bullet("is the clear winner. Instead of holding cash when the overlay de-risks, the sleeve earns a Treasuries+gold "
       "basket (TLT/GLD, IEF early). It beats the baseline on every metric — higher return (14.9% vs 12.3%), higher "
       "Sharpe (1.15 vs 1.03), SMALLER drawdown (−21% vs −26%) — and improves both in- and out-of-sample. In equity "
       "crashes (2008, 2020) Treasuries rally exactly when stocks fall, adding convexity.", lead="V1 Defensive sleeve ")
bullet("hurts: it lowers both return and Sharpe (0.94). Momentum returns concentrate in higher-volatility winners, "
       "which risk-parity weighting tilts away from. Rejected.", lead="Inverse-vol weighting ")
bullet("raise return by holding less cash / concentrating, but drawdown worsens and Sharpe barely moves (~1.05). "
       "Return dials, not free lunches — they trade away the risk differentiator.", lead="Asymmetric overlay & concentration ")
if Path("fig_bakeoff.png").exists():
    doc.add_picture("fig_bakeoff.png",width=Inches(6.3))
    cap=doc.add_paragraph(); rc=cap.add_run("Test 1 — growth of $1 across variants (SIMULATED). V1 in red.")
    rc.italic=True; rc.font.size=Pt(8); rc.font.color.rgb=LGREY

doc.add_heading("3. Test 2 — Improvements bake-off (layered on V1)",1)
body("We then tried to improve on V1 with four literature-backed enhancements, judged on out-of-sample Sharpe: "
     "risk-managed momentum (Barroso & Santa-Clara 2015), volatility targeting (Moreira & Muir 2017), a "
     "defensive (low-beta) momentum tilt (Betting-Against-Beta), and a quality-momentum composite. Note: true "
     "fundamental quality needs point-in-time data, so we used an honest price-based quality proxy (low volatility "
     "+ return stability) and label it as such.")
perf_table(W, ref_key="V1 Defensive (reference)", hi_key="V1 Defensive (reference)")
body("Findings — none of the enhancements beats V1 out-of-sample, and that is a useful result:", bold=True)
bullet("of every variant (1.14). No enhancement improves it.", lead="V1 has the highest out-of-sample Sharpe ")
bullet("are redundant with the macro overlay, which already de-grosses in stress. A second volatility-based scaler "
       "just de-risks twice — average cash rises 34%→43% and returns fall with no Sharpe gain.", lead="Risk-managed momentum & vol targeting ")
bullet("overfit: better in-sample (IS Sharpe 1.21 / 1.18) but worse out-of-sample (1.05 / 1.05) — the textbook "
       "in-sample-up / out-of-sample-down signature. Rejected.", lead="The low-beta tilt & quality proxy ")
bullet("(stacking everything) is the worst of the enhanced set — complexity actively hurt.", lead="The combined variant ")
if Path("fig_bakeoff_improvements.png").exists():
    doc.add_picture("fig_bakeoff_improvements.png",width=Inches(6.3))
    cap=doc.add_paragraph(); rc=cap.add_run("Test 2 — enhancements layered on V1 (SIMULATED). V1 reference in red.")
    rc.italic=True; rc.font.size=Pt(8); rc.font.color.rgb=LGREY

doc.add_heading("4. Recommendation",1)
body("Adopt V1 (defensive sleeve) as the flagship and stop layering.", bold=True, color=POS)
bullet("V1 is a strict improvement on the current strategy and strengthens the narrative: we rotate into convexity "
       "(Treasuries + gold), not idle cash, when the overlay de-risks.")
bullet("No further enhancement survives out-of-sample, so the parsimonious design is the right one. We tested the "
       "standard literature additions and kept the simple version by evidence, not by default — a credibility point "
       "for the pitch, and direct evidence of anti-overfitting discipline.")
bullet("Fundamental quality (ROE/profitability) is the only idea not tested 'for real'; given the price proxy already "
       "overfit, we do not expect it to change the verdict, and it would require a Bloomberg fundamentals pull.")

doc.add_heading("5. Honest caveats on V1 (to disclose)",1)
bullet("The headline alpha (+6.3%) is overstated: our FF5+Momentum model has no bond/gold factor, so the defensive "
       "sleeve's returns appear as alpha. The honest wins are the Sharpe and drawdown improvements.")
bullet("Regime dependence: the sleeve relies on bonds+gold being defensive. 2022 broke that (bonds fell with stocks), "
       "though V1's out-of-sample period — which includes 2022 — still beats the baseline. Gold diversifies this risk. "
       "Part of the historical benefit came from a multi-decade bond tailwind that will not fully repeat.")
bullet("Returns are net of trading costs but gross of fund fees and taxes; investor-net is lower.")

d=doc.add_paragraph(); rd=d.add_run("Disclaimer: Fictional pitch for the ESADE Asset Management course. Not investment advice. "
    "All performance is simulated; past performance does not indicate future results.")
rd.italic=True; rd.font.size=Pt(8); rd.font.color.rgb=LGREY

doc.save("Strategy_Refinement_Findings.docx")
print("Saved Strategy_Refinement_Findings.docx")
