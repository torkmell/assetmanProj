"""One-page fix-list for GSD2T_Presentation_v1.pptx — slide-by-slide, checkable in PowerPoint.
Output: GSD2T_Presentation_Fixlist.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor

NAVY=RGBColor(0x0B,0x1F,0x3A); GREY=RGBColor(0x55,0x55,0x55); LGREY=RGBColor(0x88,0x88,0x88)
RED=RGBColor(0xB8,0x3A,0x3A); GOLD=RGBColor(0xB8,0x86,0x0B); POS=RGBColor(0x2E,0x7D,0x5B)
BOX="☐  "  # ballot box

doc=Document(); doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(11)
def item(t, color=None, size=11, after=4, bold=False):
    p=doc.add_paragraph(); r=p.add_run(BOX); r.font.size=Pt(size)
    r2=p.add_run(t); r2.font.size=Pt(size); r2.bold=bold
    if color: r2.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); return p
def note(t,color=GREY,size=9.5):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=True; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(3); p.paragraph_format.left_indent=Pt(18); return p
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; run=c.paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(9.5)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=""; run=cs[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(9.5)
            if i==0: run.bold=True
    return t

doc.add_heading("Pitch Deck — Fix List",0)
sub=doc.add_paragraph(); r=sub.add_run("GSD²T_Presentation_v1.pptx · work through in PowerPoint and tick off")
r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GREY

doc.add_heading("CRITICAL — can't submit without these", 1)
item("Delete slides 18–51 (the template / Lorem-ipsum boilerplate). Your deck = slides 1–17.", color=RED, bold=True)
item("Fix slide 12 (Capacity): replace the garbled numbers (see the table below).", color=RED, bold=True)
item("Build out slide 15: add team names, the AI-use disclosure, and the disclaimer.", color=RED, bold=True)

doc.add_heading("Slide-by-slide", 1)
table(["Slide","Fix","Correct value / action"],[
 ["2 · Exec summary","SPY Sharpe is wrong","Change 0.65 → 0.60"],
 ["4 · Data & signals","Incomplete footnote about Sharpe","Finish the sentence or delete the footnote"],
 ["5 · The Strategy","Per-name exposure wrong","\">1% exposure per stock\" → \"~0.8% (<1%) per stock\""],
 ["6 · Portfolio constr.","Old concentration claim","Replace \"concentration — neither improved returns\" with: \"concentration looked better in-sample but REVERSED on survivorship-free data — fragile, so we hold the diversified book\""],
 ["7 · Risk framework","Team note on slide","Delete \"Double Check (is it only equity exposure…)\""],
 ["7 · Risk framework","Add coherent tail measure","Add: monthly 95% CVaR (expected shortfall) 6.7% vs market 9.6%"],
 ["8 · Back-test results","Alpha rounding","+5.3% → +5.4% (t = 4.3)"],
 ["8 · Back-test results","Chart placeholder","Replace \"5factor chart here\" with fig_factor_model.png"],
 ["8 · Back-test results","Verify the table","Sharpe column should read 1.09 / 1.16 / 1.01 / 0.60"],
 ["10 · Stress tests","Typo + team note","\"Defencive\" → \"Defensive\"; delete \"GIAN\""],
 ["11 · Robustness","Chart placeholder","Replace \"V1 line red\" with fig_bakeoff.png (V1 already red)"],
 ["12 · Capacity","STALE / garbled numbers","Soft cap $1.7–3.3B · headroom 17–33× · ~3–6% used. DELETE \"33–56×\", \"$2.3bn\", \"$3.2bn\""],
 ["12 · Capacity","Team note + add curve","Delete \"Remove / + explain\"; add fig_capacity_curve.png"],
 ["13 · Fund terms","Fee typo","\"Why not 2/12\" → \"Why not 2/20\""],
 ["15 · Thank you","Missing required content","Add team names + AI disclosure + disclaimer (see boxes below)"],
])

doc.add_heading("Global cleanup (all slides)", 1)
item("Remove every \"Source: xx\" and \"Footnote: 1. xx\" placeholder — or fill with: yfinance · Ken French Data Library · Bloomberg.")
item("Remove every leftover team note: slide 7 \"Double Check\", slide 10 \"GIAN\", slide 12 \"Remove / + explain\".")
item("Replace the two chart placeholders with the real charts: slide 8 \"5factor chart here\", slide 11 \"V1 line red\".")

doc.add_heading("Charts ready to drop in (in the assetmanProj folder)", 1)
table(["Slide","File to insert"],[
 ["7 · Risk","fig_exposure.png (exposure dial)"],
 ["8 · Back-test","fig_factor_model.png (5-factor chart) · fig_deck_equity.png (growth of $1)"],
 ["10 · Stress","fig_deck_stress.png"],
 ["11 · Robustness","fig_bakeoff.png (V1 in red)"],
 ["12 · Capacity","fig_capacity_curve.png (net Sharpe vs AUM)"],
])

doc.add_heading("Slide 15 — paste-ready content", 1)
note("Team:", color=NAVY)
item("[Name 1] — [role] · [Name 2] — [role] · [Name 3] — [role] · [Name 4] — [role]")
note("AI-use disclosure:", color=NAVY)
item("AI assistants were used for code scaffolding, back-test engineering and document drafting. All strategy "
     "decisions, data choices and results were defined and verified by the team; nothing was fabricated; every "
     "figure is reproducible from the committed code.")
note("Disclaimer:", color=NAVY)
item("Fictional pitch for the ESADE Asset Management course. Not investment advice. All performance is SIMULATED; "
     "past performance does not indicate future results. No real fund names or live numbers are used.")

doc.add_heading("Final export", 1)
item("Spec is max 15 slides: export slides 1–15 as the submission PDF (exactly 15).")
item("Keep slides 16–17 (Appendix A & B) as Q&A backup — click to them if pressed, but they're outside the 15.")
item("Remove the stray \"14\" page number on slide 16.")

d=doc.add_paragraph(); rd=d.add_run("Tip: once slides 18–51 are deleted and the boxes above are ticked, do one final "
    "read-through for any remaining \"xx\", \"insert\", or note-to-self text before exporting.")
rd.italic=True; rd.font.size=Pt(9); rd.font.color.rgb=LGREY

doc.save("GSD2T_Presentation_Fixlist.docx")
print("Saved GSD2T_Presentation_Fixlist.docx")
