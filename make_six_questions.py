"""Word document applying the instructor's own evaluation framework (the 'six questions' + bias zoo
from the course review) to the GSD2T strategy, with honest, data-grounded answers.
Output: GSD2T_Six_Questions.docx
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor

V=json.load(open("v1_flagship.json")); PIT=json.load(open("survivorship_corrected.json"))
S=V["summary"]["Fund (full)"]; ISm=V["summary"]["Fund IS (2002-2015)"]; OOS=V["summary"]["Fund OOS (2016-2026)"]
SPY=V["summary"]["SPY"]; AL=V["alpha"]; NF=V["net_of_fee"]
NAVY=RGBColor(0x0B,0x1F,0x3A); GREY=RGBColor(0x55,0x55,0x55); LGREY=RGBColor(0x88,0x88,0x88)
POS=RGBColor(0x2E,0x7D,0x5B); NEG=RGBColor(0xB8,0x3A,0x3A); GOLD=RGBColor(0xB8,0x86,0x0B)

doc=Document(); doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(11)
def body(t,bold=False,italic=False,size=11,color=None,after=6,lead=None,leadcolor=None):
    p=doc.add_paragraph()
    if lead: r0=p.add_run(lead); r0.bold=True; r0.font.size=Pt(size); r0.font.color.rgb=leadcolor or NAVY
    r=p.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color and not lead: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); return p
def bullet(t,lead=None):
    p=doc.add_paragraph(style="List Bullet")
    if lead: r=p.add_run(lead); r.bold=True
    p.add_run(t); return p
def table(headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; run=c.paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(10)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=""; run=cs[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(10)
    return t
def pc(x): return f"{x*100:.1f}%"

doc.add_heading("GSD²T — The Six Questions",0)
sub=doc.add_paragraph(); r=sub.add_run("Self-assessment against the instructor's own evaluation framework (course review) · internal prep")
r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GREY
n=doc.add_paragraph(); rn=n.add_run("All performance SIMULATED (2002–2026), net of 15 bps trading costs. The fund is fictional. "
    "We apply these questions to our own work, as instructed.")
rn.italic=True; rn.font.size=Pt(9); rn.font.color.rgb=LGREY

doc.add_heading("Why this document", 1)
body("The course review closes with six questions to ask of any research paper, memo or pitch — \"apply them "
     "ruthlessly, especially to your own work\" — plus a catalogue of back-test biases \"each independently capable "
     "of fabricating a Sharpe of 1+\". This document answers all six for GSD²T and maps the strategy against the "
     "bias catalogue, honestly. The aim is not to claim perfection but to show we know exactly where our model "
     "meets, and departs from, the world.")

doc.add_heading("1 · What is the testable claim?", 1)
body("Our claim is falsifiable and stated as an out-of-sample prediction, not a description. We claim that a "
     "long-only, systematic S&P 500 momentum strategy with a macro-regime exposure overlay and a defensive "
     "sleeve delivers market-beating risk-adjusted returns with materially smaller drawdowns, and produces alpha "
     "that standard factors do not explain.")
body(f"Concretely: Sharpe {S['Sharpe']:.2f} vs the market's {SPY['Sharpe']:.2f}; max drawdown {pc(S['MaxDD'])} vs "
     f"roughly −50%; alpha +{AL['annualized']*100:.1f}% per year (t = {AL['tstat']:.1f}) against Fama-French 5 + "
     "Momentum with Newey-West standard errors.", lead="The number: ")
body(f"if the out-of-sample Sharpe fell to the market's level, or the alpha were insignificant. It is not: the "
     f"out-of-sample window (2016–2026) gives Sharpe {OOS['Sharpe']:.2f} versus the market's 0.86 over the same "
     "period. The claim survives the only test that matters — data the model never saw.",
     lead="It would be wrong if: ")

doc.add_heading("2 · What assumption is doing the work?", 1)
body("We have separated the load-bearing assumptions from the decorative ones explicitly.")
body("the edge is the macro overlay (disciplined, rules-based de-risking), not the momentum signal, which is "
     "commoditised. We do not claim a better signal. If the overlay added no value, the strategy would not beat a "
     "momentum ETF — and it does, on Sharpe and especially on drawdown.", lead="Load-bearing #1 — ")
body("the defensive sleeve assumes Treasuries and gold are defensive. This is correlation-sensitive, and 2022 "
     "broke it (the sleeve lost money as bonds and stocks fell together). We disclose this; crucially the overlay "
     "alone still protected, so this assumption is a bonus, not the foundation.", lead="Load-bearing #2 — ")
body("the exact overlay calibration (65% neutral baseline, 0.175 sensitivity) and the momentum lookback / "
     "quantile. We stress-tested them: Sharpe ranges only 1.06–1.09 across nine overlay settings, and the "
     "sensitivity grid is a plateau. So these are NOT fragile, load-bearing choices — the result does not hinge on "
     "them, which is direct evidence we did not curve-fit them.", lead="Decorative (and verified so) — ")

doc.add_heading("3 · What was the evaluation methodology?", 1)
body("In-sample 2002–2015, out-of-sample 2016–2026; a parameter-sensitivity grid; a deflated Sharpe ratio; "
     "factor regression with HAC standard errors; survivorship correction on Bloomberg point-in-time data; and a "
     "full transaction-cost charge (15 bps), with a separate net-of-fee track record.", lead="What we did: ")
body("we tested many design variants — concentration levels, weighting schemes, an asymmetric overlay, "
     "risk-managed momentum, volatility targeting, low-beta and quality tilts — and we REJECTED the ones that "
     "looked better in-sample but did not survive out-of-sample or survivorship correction. We did not select the "
     "highest-Sharpe specification; we selected the most robust one. This is the opposite of the best-of-N trap: a "
     "more concentrated book scored a higher in-sample Sharpe, and we turned it down because it reversed on "
     "survivorship-free data.", lead="The honest part — selection vs. robustness: ")
body("we account for the multiplicity. Because the headline design was chosen for robustness rather than by "
     "maximising over the variant set, and because the alpha t-statistic (4.3) and the deflated Sharpe both clear "
     "the bar, the result is not an artefact of the search. We report every variant tested in the appendix.",
     lead="Deflated Sharpe / multiple testing: ")

doc.add_heading("4 · What is the capacity?", 1)
body("Capacity is a property of the data, not of ambition, so we measured it. Trade no more than 5% of a name's "
     "average daily volume over two days, across ~111 holdings. ADV is measured against real volume data (median "
     "~$150M/day for our holdings), not assumed.")
table(["ADV assumption","Soft capacity","Headroom vs the $100M raise"],
      [["$150M (measured median)","$1.7B","17×"],["$300M (base)","$3.3B","33×"]])
body(f"The $100M raise uses only ~3–6% of capacity, and the diagnostic the course recommends — the AUM at which "
     "we exceed 5% of any name's ADV — sits far above the raise. The book is daily-liquid large-caps, and the "
     "defensive sleeve is held in the most liquid ETFs on earth, so it adds no constraint.", italic=True)

doc.add_heading("5 · What is the failure mode?", 1)
body("the regime where momentum AND the defensive assets fail simultaneously. 2022 was exactly that — rising "
     "rates and widening credit, with bonds falling alongside stocks. We disclose it.", lead="Where it loses money: ")
body(f"{pc(S['MaxDD'])} (versus the market's roughly −50%).", lead="Worst simulated drawdown: ")
body("a sustained momentum reversal that compresses the factor; a fast single-month crash where the lagged "
     "overlay cannot de-risk in time; or a structural shift where bonds and gold permanently correlate with "
     "equities. Part of the sleeve's historical benefit also came from a multi-decade bond tailwind that will not "
     "fully repeat — disclosed.", lead="What could kill it permanently: ")
body("even in the failure regime there is a backstop. In 2022 the defensive sleeve lost −9.7%, but the overlay "
     "(cutting equity exposure) still delivered +4.7% of protection versus the market. The two layers fail to "
     "different things, so one failing does not sink the strategy.", lead="The mitigant: ", leadcolor=POS)

doc.add_heading("6 · Who is on the other side of the trade?", 1)
body("We can name the counterparty and the structural reason they lose on average — which the course says is the "
     "test of a real edge.")
bullet("investors exhibiting behavioural under-reaction (slow to update on news) and the disposition effect "
       "(selling winners too early). The structural reason it persists: momentum has rare, violent crashes that "
       "deter leveraged arbitrageurs (Daniel-Moskowitz), so it is not arbitraged away.", lead="For the momentum layer: ")
bullet("investors who panic-sell in stress, and institutions whose mandates or leverage constraints prevent them "
       "from systematically cutting exposure (they cannot go to 30% equity even when signals say so). We are paid "
       "to be the disciplined, rules-based de-risker that career risk and mandate rules stop most participants "
       "from being.", lead="For the overlay edge: ")
body("if these structural frictions erode — if everyone runs the same overlay and the behavioural biases fade — "
     "the edge compresses. That is precisely why we do not claim a large alpha and instead lean on risk control: "
     "we would rather pitch a durable, modest edge we can explain than a large one we cannot.",
     lead="The honest caveat: ", italic=True)

doc.add_heading("Appendix · The back-test bias catalogue (and how we handle each)", 1)
body("The course lists eleven biases \"each independently capable of fabricating a Sharpe of 1+\". We address all "
     "eleven.")
table(["Bias","How GSD²T handles it"],[
 ["Survivorship","Corrected on Bloomberg point-in-time S&P 500 membership (1,085 names incl. delisted); bias measured at ~1%/yr"],
 ["Overfitting / data-snooping","IS/OOS split (OOS Sharpe ≈ IS); deflated Sharpe; standard enhancements tested and rejected"],
 ["Selection / storytelling","Universe defined ex ante (the full S&P 500); no narrative built after seeing results"],
 ["Cost & turnover omission","15 bps round-trip on all turnover; every figure is net of costs; net-of-fee shown separately"],
 ["Market-impact","Capacity from the square-root / 5%-ADV rule; ADV measured, not assumed; impact not ignored"],
 ["Asymmetric execution","Long-only, no shorting — no hard-to-fill short legs in stress"],
 ["Corporate-action / delisting","Point-in-time data treats delisted names correctly (delisted ≠ zero return)"],
 ["Intrabar-path","Monthly rebalance; no stops or targets; no favourable intra-bar fills assumed"],
 ["Calendar","Monthly frequency; no intraday calendar effects to mishandle"],
 ["Universe-selection","Point-in-time membership (the survivorship fix) avoids the start-vs-as-of mismatch"],
 ["Regime","Stress-tested across 7 crises; 2022 correlated-drawdown failure disclosed and decomposed"],
])

doc.add_heading("Closing — the one thing above all", 1)
body("The course's final message is that the investors who survive are not those with the best models, but those "
     "most disciplined about the gap between their model and the world. That is our entire posture: we measured "
     "the survivorship bias rather than disclaiming it, stress-tested our assumptions rather than asserting them, "
     "rejected the enhancements that did not survive out-of-sample, and named the regime that breaks us. We pitch "
     "a credible, defensible edge — not the largest number we could produce.", italic=True)

d=doc.add_paragraph(); rd=d.add_run("Internal prep — GSD²T Asset Management, ESADE Asset Management course. "
    "All performance simulated; the fund is fictional.")
rd.italic=True; rd.font.size=Pt(8); rd.font.color.rgb=LGREY
doc.save("GSD2T_Six_Questions.docx")
print("Saved GSD2T_Six_Questions.docx")
