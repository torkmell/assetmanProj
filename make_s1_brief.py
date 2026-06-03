"""Generate a high-level Word brief for Strategy 1 (Macro-Overlay Tech) to share with the team."""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Pull live headline numbers from the backtest output so the doc stays accurate
G = json.load(open("gsd2t_full.json"))
S = G["summary"]["Fund (full window)"]
SPY = G["summary"]["SPY"]
FR = G["factor_regression"]
hold = G["holdings_latest"]["weights"]
n_now = len(hold)
gross_now = sum(x["weight"] for x in hold)

doc = Document()

# ---- base styling ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def table(headers, rows, style="Light Grid Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = style
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(hd); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t

# ============================================================ TITLE
title = doc.add_heading("GSD2T — Strategy 1: Macro-Overlay Tech", level=0)
sub = doc.add_paragraph()
r = sub.add_run("Team brief · ESADE Asset Management · The pitched fund (S1)")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
note = doc.add_paragraph()
rn = note.add_run("All performance is SIMULATED on historical data (2002–2026). The fund is fictional.")
rn.italic = True; rn.font.size = Pt(9); rn.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ============================================================ 1. THE IDEA
h("1. The strategy in one paragraph", 1)
body("S1 buys the strongest-performing US technology stocks (a momentum signal) and then "
     "uses a simple macroeconomic “risk dial” to decide how much of the portfolio to hold "
     "in those stocks versus cash. In calm, healthy markets the dial turns exposure up toward "
     "100%; when stress signals flash (rising volatility, widening credit spreads, rising rates, "
     "weakening trend) it turns exposure down toward a 30% floor and parks the rest in cash. "
     "The result is tech-like upside with materially smaller drawdowns.")

# ============================================================ 2. TWO LAYERS
h("2. How it works — two layers", 1)
body("The strategy is built from two independent layers:", bold=True)
bullet("Layer 1 – Stock selection: pick which tech stocks to own (momentum).")
bullet("Layer 2 – Macro overlay: decide how much to own overall (the risk dial).")
body("Final position in each stock  =  (equal weight within the chosen stocks)  ×  (macro exposure level).",
     italic=True)

# ============================================================ 3. LAYER 1
h("3. Layer 1 — Stock selection (the momentum factor)", 1)
table(
    ["Item", "Detail"],
    [
        ["Universe", "S&P 500 ∩ Information Technology — 73 stocks"],
        ["Factor", "12-minus-1 month price momentum"],
        ["How computed", "Total return over the last 12 months, excluding the most recent month "
                          "(skip-a-month avoids short-term reversal)"],
        ["Ranking", "Each stock z-scored against the universe each month"],
        ["Selection", "Hold the top quartile (best ~25%) ≈ 18–19 stocks"],
        ["Weighting", "Equal weight within the selected stocks"],
    ],
)
body("Why momentum: winners tend to keep winning over 3–12 month horizons, a well-documented "
     "effect (Jegadeesh & Titman, 1993). Equal-weighting keeps the book diversified and prevents "
     "any single name from dominating.", italic=True)

# ============================================================ 4. LAYER 2
h("4. Layer 2 — The macro overlay (4 factors)", 1)
body("Four economically-motivated signals are each standardised (z-scored over a rolling 5-year "
     "window) and averaged into one composite “risk score”. Three are stress signals (turn risk "
     "down) and one is a trend signal (turn risk up).")
table(
    ["Macro factor", "What it captures", "Data used", "How computed", "Direction"],
    [
        ["VIX", "Equity fear / volatility", "CBOE VIX index", "z-score of VIX level", "High → de-risk"],
        ["Credit", "Credit stress", "IEF vs HYG ETFs", "z-score of 12-month change in the Treasury-minus-high-yield gap", "Widening → de-risk"],
        ["Yield", "Rate tightening", "10-year Treasury yield", "z-score of 12-month change in the 10Y yield", "Rising → de-risk"],
        ["Trend", "Market momentum", "S&P 500 index", "z-score of the 12-1 month S&P 500 return", "Up → risk-on"],
    ],
)
body("Composite risk score = average of the four factors (each equally weighted), lagged one month "
     "so we never use data we would not have had in real time.", italic=True)

h("From risk score to exposure", 2)
body("The composite maps to the share of the portfolio held in stocks (the rest sits in cash, "
     "earning the risk-free rate):")
table(
    ["Risk score", "Stock exposure", "Cash"],
    [["+2 (calm / risk-on)", "100% (capped)", "0%"],
     ["0 (neutral)", "65%", "35%"],
     ["−2 (stress)", "30% (floor)", "70%"]],
)

# ============================================================ 5. WEIGHTING
h("5. Putting it together — weighting", 1)
body("Weight per stock  =  (1 / number of stocks held)  ×  macro exposure", bold=True)
body(f"Live example (latest month): {n_now} stocks held, macro exposure at {gross_now*100:.0f}%, "
     f"so each stock = {gross_now/n_now*100:.2f}% and {100-gross_now*100:.0f}% sits in cash. "
     f"At full 100% exposure each stock would be ~{100/n_now:.1f}%.")
bullet("Equal-weight → built-in single-name cap (~5%) and a minimum of ~18 holdings.")
bullet("No leverage, no shorting; weights in the stock sleeve sum to the macro exposure level.")
bullet("Rebalanced monthly; 15 bps round-trip transaction cost assumed.")

# ============================================================ 6. DATA
h("6. Data sources", 1)
table(
    ["Data", "Source", "Used for"],
    [
        ["Stock & ETF prices (monthly, dividend-adjusted)", "yfinance", "Momentum signal, returns"],
        ["VIX index", "yfinance (^VIX)", "Macro overlay – volatility"],
        ["IEF, HYG bond ETFs", "yfinance", "Macro overlay – credit spread proxy"],
        ["10-year Treasury yield (^TNX)", "yfinance", "Macro overlay – rates"],
        ["S&P 500 index (^GSPC)", "yfinance", "Macro overlay – trend"],
        ["Fama-French 5 factors + Momentum, risk-free rate", "Ken French data library", "Risk-free rate & alpha test"],
    ],
)
body("Note: credit and rates currently use ETF/price proxies; an institutional build would upgrade "
     "these to true high-yield spread and Treasury series (e.g. via Bloomberg).", italic=True)

# ============================================================ 7. RESULTS
h("7. Headline results (simulated, 2002–2026, net of costs)", 1)
table(
    ["Metric", "S1 Fund", "S&P 500 (SPY)"],
    [
        ["CAGR", f"{S['CAGR']*100:.1f}%", f"{SPY['CAGR']*100:.1f}%"],
        ["Sharpe ratio", f"{S['Sharpe']:.2f}", f"{SPY['Sharpe']:.2f}"],
        ["Max drawdown", f"{S['MaxDD']*100:.0f}%", f"{SPY['MaxDD']*100:.0f}%"],
        ["Alpha vs FF5+Momentum", f"+{FR['alpha_annualized']*100:.1f}% (t={FR['alpha_tstat']:.1f})", "—"],
    ],
)
body("In plain terms: comparable-to-better return than the S&P 500 at roughly half the drawdown, "
     "with statistically significant alpha after accounting for standard risk factors.", italic=True)

# ============================================================ FOOTER
doc.add_paragraph()
disc = doc.add_paragraph()
rd = disc.add_run("Disclaimer: Fictional pitch for the ESADE Asset Management course. Not investment "
                  "advice. All performance is simulated; past performance does not indicate future results. "
                  "Universe uses current S&P 500 constituents (survivorship bias disclosed).")
rd.italic = True; rd.font.size = Pt(8); rd.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

out = "S1_Strategy_Overview.docx"
doc.save(out)
print(f"Saved {out}")
