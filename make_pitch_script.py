"""Run-of-show + speaker script for the GSD2T live pitch.
One page-ish per block: what each presenter SAYS on every slide, the numbers to hit,
and the hand-offs. British English, no em/en dashes (house style).
Output: GSD2T_Pitch_Script.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY=RGBColor(0x0B,0x1F,0x3A); GOLD=RGBColor(0xB8,0x86,0x0B); GREY=RGBColor(0x55,0x55,0x55)
LGREY=RGBColor(0x88,0x88,0x88); POS=RGBColor(0x2E,0x7D,0x5B); BLACK=RGBColor(0x1A,0x1A,0x1A)

doc=Document()
n=doc.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(11); n.font.color.rgb=BLACK

def H(txt,size=22,color=NAVY,after=2,before=0):
    p=doc.add_paragraph(); r=p.add_run(txt); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); return p
def sub(txt,size=11,color=GREY,italic=True,after=8):
    p=doc.add_paragraph(); r=p.add_run(txt); r.italic=italic; r.font.size=Pt(size); r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(after); return p
def slidehead(num,title,speaker,time):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(f"Slide {num}  ·  {title}"); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=NAVY
    r2=p.add_run(f"     [{speaker} · {time}]"); r2.bold=True; r2.font.size=Pt(10.5); r2.font.color.rgb=GOLD
def say(txt):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.left_indent=Pt(10)
    r0=p.add_run("Say:  "); r0.bold=True; r0.font.size=Pt(10.5); r0.font.color.rgb=GREY
    r=p.add_run(txt); r.font.size=Pt(11); r.italic=False
def hit(txt):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Pt(10)
    r0=p.add_run("Numbers to hit:  "); r0.bold=True; r0.font.size=Pt(10); r0.font.color.rgb=POS
    r=p.add_run(txt); r.font.size=Pt(10); r.font.color.rgb=BLACK
def handoff(txt):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Pt(10)
    r0=p.add_run("Hand off:  "); r0.bold=True; r0.font.size=Pt(10); r0.font.color.rgb=GOLD
    r=p.add_run(txt); r.italic=True; r.font.size=Pt(10); r.font.color.rgb=GREY
def table(headers,rows,widths=None,fs=10):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(fs)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=""; r=cs[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(fs)
            if i==0: r.bold=True
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    return t

# ===================== HEADER =====================
H("GSD²T Asset Management",26,NAVY,after=0)
H("Pitch Run-of-Show & Speaker Script",15,GOLD,after=2)
sub("15-minute pitch + 5-minute Q&A  ·  ESADE Asset Management  ·  live pitch 10 June 2026")

p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8)
r=p.add_run("How to use this: ").bold=True
p.add_run("each presenter owns the block in their colour. Read the ").italic=True
rr=p.add_run("Say"); rr.italic=True; rr.bold=True
p.add_run(" lines as a guide, not a recital, then make sure the ").italic=True
rr=p.add_run("Numbers to hit"); rr.italic=True; rr.bold=True
p.add_run(" actually leave your mouth. Aim to land at 13 to 14 minutes so the 5-minute Q&A is never squeezed.").italic=True

# ===================== AT A GLANCE =====================
H("At a glance",15,NAVY,after=4,before=6)
table(["Speaker","Slides","Owns","Scored area"],[
 ["1 — the hook","1, 2, 3","Title, Executive Summary, The Edge","Edge & strategy (25)"],
 ["2 — how it works","4, 5, 6","Data, Strategy, Portfolio","Edge & strategy (25)"],
 ["3 — risk & proof","7, 8","Risk Framework, Back-test","Risk framework (15)"],
 ["4 — honesty","9, 10, 11","Honesty, Stress, Robustness","Back-test honesty (25)"],
 ["5 — the business","12, 13, 14, 15","Capacity, Terms, Why Now, Close","Fund terms (10) + leads Q&A"],
],widths=[1.5,1.0,2.5,1.7],fs=9.5)
sub("Four-person team: merge Speakers 3 and 4 into one person on slides 7 to 11. Put your calmest "
    "presenter on the honesty block (Speaker 4): that is where the hostile questions land, and composure is worth 20 points.",
    size=9.5,after=6)

# ===================== SCRIPT =====================
H("Slide-by-slide script",15,NAVY,after=4,before=6)

slidehead(1,"Title","Speaker 1","~15 sec")
say("Good morning. We are GSD²T Asset Management, and we are raising 100 million dollars for a long-only, "
    "systematic S&P 500 strategy. It combines momentum stock selection, a macro overlay that manages risk, and a "
    "defensive sleeve of Treasuries and gold. Over the next fifteen minutes we will show you why it works, prove it "
    "honestly, and explain how we make money from it.")

slidehead(2,"Executive Summary","Speaker 1","~90 sec")
say("Here is the whole pitch on one slide. Simulated over twenty-four years, the strategy compounds at 13.1 per cent "
    "net of fees, against the market's 10. It does that at a higher Sharpe ratio, 1.09 versus 0.60, and with less than "
    "half the drawdown, minus 21 versus minus 51. And the outperformance is not luck: after stripping out the standard "
    "equity factors, 5.4 per cent of annual alpha remains, with a t-statistic of 4.3, and it holds out of sample. "
    "The rest of the deck is us proving each of those claims and then trying to break them.")
hit("13.1% net vs 10.0% · Sharpe 1.09 vs 0.60 · drawdown -21% vs -51% · alpha +5.4% (t = 4.3)")

slidehead(3,"The Edge","Speaker 1","~90 sec")
say("So what is the edge? It is not a cleverer momentum signal; those are commoditised. Our edge is disciplined, "
    "rules-based de-risking. Investors overpay in calm markets and panic-sell in stress. Momentum persists because of "
    "under-reaction, and drawdowns persist because risk premia move over time. The reason it does not get arbitraged "
    "away is behavioural and structural: mandates, career risk and leverage limits stop most managers from cutting "
    "exposure when their own signals tell them to. We have no such constraint, because our process is mechanical.")
handoff("“To show you how the machine works, I will hand over to [name].”")

slidehead(4,"Data & Signals","Speaker 2","~60 sec")
say("Thank you. The engine has two layers. Layer one is selection: twelve-minus-one month price momentum across the "
    "full S&P 500, ranked cross-sectionally each month, and we equal-weight the top quartile, about 125 names. Layer "
    "two is the macro overlay: four factors, the VIX for equity fear, a credit spread, the change in ten-year yields, "
    "and market trend. Every input uses only data available at decision time, lagged by at least a month, and we trade "
    "the following month. There is no look-ahead anywhere in the pipeline.")
hit("12-1 momentum · top quartile (~125 names) · 4 factors · all lagged ≥ 1 month")

slidehead(5,"The Strategy","Speaker 2","~90 sec")
say("Putting it together in three steps. One, we select: the strongest momentum names across all eleven sectors, "
    "equal-weighted. Two, we overlay: the four-factor risk dial scales our equity exposure between 30 and 100 per cent. "
    "Calm markets, fully invested; stress, we de-risk hard. Three, and this is our key refinement, the de-risked "
    "portion does not sit in cash. It earns in Treasuries and gold, which tend to rise when equities fall. So our "
    "risk-off bucket is productive: it pays us while we wait and it cushions us in a crisis.")
hit("exposure 30–100% · de-risked portion in Treasuries + gold, not cash")

slidehead(6,"Portfolio Construction","Speaker 2","~60 sec")
say("On construction, we are genuinely diversified. Today's book is 126 names across all eleven sectors, roughly one "
    "per cent each, so no single name can hurt us. We did test concentration, holding fewer, higher-conviction names. "
    "It looked better in the backtest, but that edge reverses once we remove survivorship bias, so it is an artefact, "
    "not a real effect. We hold the diversified book on purpose.")
hit("126 names · 11 sectors · ~1% each · concentration tested and rejected as a survivorship artefact")
handoff("“[Name] will now take you through the risk framework and the results.”")

slidehead(7,"Risk Framework","Speaker 3","~90 sec")
say("Thanks. This chart is the heart of the strategy: our risk dial in action over twenty-four years. The line is our "
    "equity exposure. In calm periods we are near fully invested. In every major stress event, 2008, 2020 and 2022, it "
    "drops to our 30 per cent equity floor, and the rest moves into the Treasuries and gold sleeve. Everything else on "
    "this slide follows from that line: a beta of 0.69, volatility of 11.9 against the market's 15, a maximum drawdown "
    "of minus 21 versus minus 51, and a 95 per cent CVaR, our expected loss in the worst months, of 6.7 per cent "
    "against the market's 9.6. We win on risk, the denominator, not by chasing return.")
hit("beta 0.69 · vol 11.9% vs 15.1% · maxDD -21% vs -51% · CVaR 6.7% vs 9.6% · equity floor 30%")

slidehead(8,"Back-test Results","Speaker 3","~90 sec")
say("And here is the performance. One dollar becomes about thirty over the period, versus ten for the market: 14.9 per "
    "cent compound growth against 10. The table shows it holds both in-sample and out-of-sample, with the out-of-sample "
    "drawdown actually smaller. The chart on the right is the one that matters for this audience: when we regress our "
    "returns on the five Fama-French factors plus momentum, 5.4 per cent of annual alpha remains after all of them. "
    "This is not disguised market beta.")
hit("14.9% vs 10.0% CAGR · IS Sharpe 1.16 / OOS 1.01 · alpha +5.4% after 6 factors")
handoff("“To show you we are not fooling ourselves, [name] will walk through our honesty checks.”")

slidehead(9,"Back-testing Honesty","Speaker 4","~90 sec")
say("Thank you. Every backtest can lie, so we went looking for the ways ours could. Survivorship bias: we measured it "
    "directly on Bloomberg point-in-time data, including delisted names, and it is about one per cent a year, "
    "quantified, not assumed. Costs: fifteen basis points round-trip on all turnover, and every figure is net. "
    "Look-ahead: every signal is lagged, with trailing-only windows. Overfitting: we trained on 2002 to 2015 and tested "
    "on 2016 onward, and the Sharpe holds, 1.01 out of sample versus 1.16 in. It steps down modestly, but it does not "
    "collapse. We are showing you the honest number.")
hit("survivorship ~1%/yr (point-in-time Bloomberg, 1,085 names) · 15bps net · all lagged · OOS 1.01 vs IS 1.16")

slidehead(10,"Stress Tests","Speaker 4","~60 sec")
say("This is the same idea event by event. In every major crisis we fall far less than the market: a 38-point gap in "
    "the dot-com bust and 36 in the financial crisis. We deliberately include 2022, the one regime where bonds and "
    "stocks fell together and our sleeve helped least. We disclose it because it is a real risk, and even so, our "
    "out-of-sample period, which includes 2022, still beats the market.")
hit("Dot-com gap +38 · GFC +36 · 2022 disclosed (bonds and stocks fell together)")

slidehead(11,"Robustness","Speaker 4","~75 sec")
say("Finally on the research, we tried to break the strategy. The heatmap shows the Sharpe ratio across fifteen "
    "different parameter settings, and it sits in a tight band, 0.93 to 1.00. That is a plateau, not a lucky spike, "
    "which means we did not tune to one magic setting. We also tested the popular enhancements, vol targeting, quality "
    "tilts, risk-managed momentum, and none of them beat the simple design out of sample. We kept it simple because the "
    "evidence told us to.")
hit("Sharpe 0.93 to 1.00 across 15 settings (a plateau) · enhancements tested, none beat the simple design OOS")
handoff("“[Name] will now cover capacity, terms, and why now.”")

slidehead(12,"Capacity","Speaker 5","~60 sec")
say("Thank you. Capacity is where ambition often outruns reality, so we measured it rather than assuming it. Based on "
    "real average daily volume, taking five per cent over two days across our holdings, our soft capacity is between "
    "1.7 and 3.3 billion dollars. The 100 million we are raising is three to six per cent of that, so we have "
    "seventeen to thirty-three times headroom. The full amount deploys immediately in daily-liquid large caps. We never "
    "come close to moving our own market.")
hit("soft cap $1.7B to $3.3B · raise is 3 to 6% of capacity · 17 to 33x headroom")

slidehead(13,"Fund Structure & Terms","Speaker 5","~60 sec")
say("On terms, we align with our liquidity. A one per cent management fee, and a fifteen per cent performance fee "
    "charged only on return above the S&P 500, not on market beta you can buy for five basis points. There is a "
    "relative high-water mark, and monthly liquidity with no lock-up. We are daily-liquid and partly factor-replicable, "
    "so we do not charge two-and-twenty. Net of all fees, the investor still earns 13.1 per cent against the market's 10.")
hit("1% mgmt · 15% over the S&P · relative high-water mark · monthly liquidity · net 13.1% vs 10.0%")

slidehead(14,"Why Now","Speaker 5","~45 sec")
say("Why now? Three reasons. Regime uncertainty is elevated, so a strategy built to de-risk systematically is timely. "
    "Defensive assets pay a real yield again, so our Treasuries and gold sleeve earns while it waits, unlike the "
    "zero-rate decade. And with raw signals commoditised, the durable edge has moved to risk management and execution, "
    "which is exactly our strength.")

slidehead(15,"Thank You","Speaker 5","~15 sec")
say("That is GSD²T: market-beating compounding, honestly measured, at half the risk. Thank you. We would welcome "
    "your questions.")

# ===================== Q&A PLAYBOOK =====================
H("Q&A playbook  (the other 5 minutes)",15,NAVY,after=4,before=10)
sub("Speaker 5 receives every question, repeats it back in one line, then routes it to the owner. Visible coordination "
    "is itself scored. Concede what is true before you defend: measured honesty wins more points than a perfect answer.",
    size=10,after=6)
table(["Likely question","Who answers","Where it is backed"],[
 ["Why not concentrate in 25 high-conviction names?","Speaker 2 / 4","Appendix A (slide 16) — reverses survivorship-free"],
 ["Isn't your result sensitive to costs or the overlay?","Speaker 4","Appendix B (slide 17) — holds at 6x cost; overlay is a plateau"],
 ["Why no machine learning?","Speaker 2","Q&A doc — signal is commoditised; edge is in risk control"],
 ["Why long-only, not long-short?","Speaker 3","Q&A doc — borrow cost, capacity, and crash risk on the short book"],
 ["Isn't this just momentum plus market beta?","Speaker 3","Slide 8 — +5.4% alpha after FF5 + momentum"],
 ["Why hold less equity than a normal momentum fund?","Speaker 2 / 3","The sleeve is productive, so de-grossing costs us nothing"],
 ["Your OOS Sharpe is lower than in-sample — is it decaying?","Speaker 4","Concede: 1.16 to 1.01 is a modest step-down, not a collapse"],
],widths=[3.0,1.3,2.7],fs=9.5)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8)
r=p.add_run("Full preparation: ").bold=True
p.add_run("the 39-question simulation in ")
rr=p.add_run("GSD2T_QA_Simulation.docx"); rr.font.name="Consolas"; rr.font.size=Pt(10)
p.add_run(" covers every angle above and more. Rehearse the first answer out loud until it is twenty seconds, not sixty.")

doc.save("GSD2T_Pitch_Script.docx")
print("Saved GSD2T_Pitch_Script.docx")
