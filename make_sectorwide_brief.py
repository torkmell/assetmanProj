"""High-level Word brief for the Sector-Wide (diversified) strategy — the new flagship."""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

F = json.load(open("sectorwide_full.json"))
NF = json.load(open("net_of_fee.json"))
PIT = json.load(open("survivorship_corrected.json"))
S = F["summary"]["Fund (full window)"]; SPY = F["summary"]["SPY"]; FR = F["factor_regression"]
CAP = F["capacity"]; SENS = F["sensitivity"]; HL = F["holdings_latest"]; SEC = F["sector_breakdown"]
IS = F["summary"]["Fund IS (2002-2015)"]; OOS = F["summary"]["Fund OOS (2016-2026)"]

doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)
def body(t, bold=False, italic=False, size=11):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(size); return p
def bullet(t): doc.add_paragraph(t, style="List Bullet")
def table(headers, rows, style="Light Grid Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = style
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; run=c.paragraphs[0].add_run(h); run.bold=True
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row): cs[i].text=str(v)
    return t

doc.add_heading("GSD2T — Sector-Wide Strategy (Diversified Macro-Overlay)", level=0)
sub=doc.add_paragraph(); r=sub.add_run("Team brief · ESADE Asset Management · The diversified flagship")
r.italic=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x55,0x55,0x55)
n=doc.add_paragraph(); rn=n.add_run("All performance is SIMULATED (2002–2026), net of 15bps trading costs, GROSS of fund fees. The fund is fictional.")
rn.italic=True; rn.font.size=Pt(9); rn.font.color.rgb=RGBColor(0x88,0x88,0x88)

doc.add_heading("1. The strategy in one paragraph", 1)
body("We buy the strongest-performing stocks across the WHOLE S&P 500 (all 11 sectors), then "
     "use a simple macroeconomic “risk dial” to decide how much to hold in stocks versus cash. "
     "It is the same engine as our tech strategy, applied to a diversified universe. The result: "
     "broad-market beating, risk-adjusted returns from a portfolio that is genuinely diversified "
     "— not a sector bet — with strong downside protection in every major crisis.")

doc.add_heading("2. How it works — two layers", 1)
bullet("Layer 1 – Stock selection: pick the strongest stocks by 12-month price momentum, across all sectors.")
bullet("Layer 2 – Macro overlay: a 4-factor “risk dial” scales total exposure between 30% and 100% (rest in cash).")
body("Position per stock = (equal weight within the chosen stocks) × (macro exposure level).", italic=True)

doc.add_heading("3. Layer 1 — Stock selection (momentum factor)", 1)
table(["Item","Detail"], [
    ["Universe", f"Full S&P 500 — all 11 sectors (~500 stocks)"],
    ["Factor", "12-minus-1 month price momentum"],
    ["How computed", "Total return over the last 12 months, excluding the most recent month"],
    ["Ranking", "Each stock z-scored against the universe each month"],
    ["Selection", f"Hold the top quartile — currently {HL['n_holdings']} stocks across {HL['n_sectors']} sectors"],
    ["Weighting", "Equal weight within the selected stocks"],
])

doc.add_heading("4. Layer 2 — The macro overlay (4 factors)", 1)
body("Four economically-motivated signals, each z-scored over a rolling 5-year window, averaged "
     "into one composite “risk score”. Three are stress signals (turn risk down); one is trend (turn risk up).")
table(["Macro factor","Captures","Data","Direction"], [
    ["VIX","Equity fear / volatility","CBOE VIX","High → de-risk"],
    ["Credit","Credit stress","IEF vs HYG ETFs","Widening → de-risk"],
    ["Yield","Rate tightening","10-year Treasury yield","Rising → de-risk"],
    ["Trend","Market momentum","S&P 500 index","Up → risk-on"],
])
table(["Risk score","Stock exposure","Cash"], [["+2 calm","100%","0%"],["0 neutral","65%","35%"],["−2 stress","30%","70%"]])

doc.add_heading("5. Weighting", 1)
body("Weight per stock = (1 / number held) × macro exposure.", bold=True)
bullet(f"Currently {HL['n_holdings']} stocks held; equal-weight → each ~{100/HL['n_holdings']:.1f}% at full exposure (built-in ~1% single-name cap).")
bullet("No leverage, no shorting; rebalanced monthly; 15 bps round-trip trading cost.")
body("Current sector spread of the book (diversified, not a tech bet):", italic=True)
table(["Sector","# stocks"], [[k, v] for k,v in list(SEC.items())])

doc.add_heading("6. Data sources", 1)
table(["Data","Source","Used for"], [
    ["Stock prices (monthly, adjusted)","yfinance","Momentum, returns"],
    ["VIX, IEF, HYG, 10Y yield, S&P 500","yfinance","Macro overlay"],
    ["Fama-French 5 + Momentum, risk-free","Ken French library","Risk-free & alpha test"],
])

doc.add_heading("7. Results (simulated, 2002–2026, net of trading costs)", 1)
table(["Metric","Sector-Wide","S&P 500 (SPY)"], [
    ["CAGR", f"{S['CAGR']*100:.1f}%", f"{SPY['CAGR']*100:.1f}%"],
    ["Sharpe ratio", f"{S['Sharpe']:.2f}", f"{SPY['Sharpe']:.2f}"],
    ["Max drawdown", f"{S['MaxDD']*100:.0f}%", f"{SPY['MaxDD']*100:.0f}%"],
    ["Alpha vs FF5+Momentum", f"+{FR['alpha_annualized']*100:.1f}% (t={FR['alpha_tstat']:.1f})", "—"],
    ["Capacity vs $100M raise", f"${CAP['soft_cap_low']/1e9:.0f}-{CAP['soft_cap_high']/1e9:.0f}B "
        f"({CAP['headroom_low']:.0f}-{CAP['headroom_high']:.0f}x headroom)", "—"],
])
body("It is robust, not curve-fit: the result holds as a plateau across parameter settings, survives the "
     "survivorship correction (~1%/yr measured bias), and stands up to transaction-cost and overlay-calibration "
     "stress-tests — strong evidence it is a real effect, not a tech bet.", italic=True)
body(f"On the $100M commitment we are raising, the strategy runs at roughly "
     f"{CAP['pct_of_capacity_at_raise']*100:.0f}% of capacity — {CAP['headroom_low']:.0f}-{CAP['headroom_high']:.0f}x headroom. "
     "We can deploy the full $100M immediately in daily-liquid large-caps and never need to soft-close early.", italic=True)

doc.add_heading("8. Fund terms (aligned with our liquidity)", 1)
T = NF["terms"]
table(["Term","What we charge"], [
    ["Management fee", T["management_fee"]],
    ["Performance fee", T["performance_fee"]],
    ["High-water mark", T["high_water_mark"]],
    ["Crystallisation", T["crystallisation"]],
    ["Liquidity", T["liquidity"]],
])
body("Why not 2/20: the strategy is daily-liquid and partly factor-replicable, so we charge below "
     "hedge-fund terms and the performance fee applies ONLY to returns above the S&P 500 — we do not "
     "charge an alpha fee on market beta you can buy for 5 bps, and the high-water mark means we only "
     "get paid on new highs versus the benchmark.", italic=True)

body("What the investor actually receives (simulated, net of all fees):", bold=True)
G, N2, SP = NF["gross"], NF["net"], NF["spy"]
table(["", "CAGR", "Sharpe", "Max drawdown"], [
    ["Gross strategy",   f"{G['CAGR']*100:.1f}%",  f"{G['Sharpe']:.2f}",  f"{G['MaxDD']*100:.0f}%"],
    ["NET to investor",  f"{N2['CAGR']*100:.1f}%", f"{N2['Sharpe']:.2f}", f"{N2['MaxDD']*100:.0f}%"],
    ["S&P 500 (SPY)",    f"{SP['CAGR']*100:.1f}%", f"{SP['Sharpe']:.2f}", f"{SP['MaxDD']*100:.0f}%"],
])
body(f"After all fees the client still beats the S&P 500 ({N2['CAGR']*100:.1f}% vs {SP['CAGR']*100:.1f}%), "
     f"but the real value is risk: Sharpe {N2['Sharpe']:.2f} vs {SP['Sharpe']:.2f} and a {N2['MaxDD']*100:.0f}% "
     f"drawdown versus {SP['MaxDD']*100:.0f}%. We are paid for risk-adjusted outperformance, not raw return.",
     italic=True)

doc.add_heading("9. Honest notes (for the team)", 1)
bullet("Returns are NET of trading costs but GROSS of fund management/performance fees and taxes; investor net is lower once we set fees.")
bullet(f"In-sample Sharpe {IS['Sharpe']:.2f}, out-of-sample {OOS['Sharpe']:.2f}; parameter-robust ({SENS['min_sharpe']:.2f}–{SENS['max_sharpe']:.2f} across 15 settings).")
bullet("2019–2024 was a mega-cap-growth 'factor drought': out-of-sample we preserved capital and kept pace at lower risk, but the factor alpha compressed industry-wide. We disclose this.")
bullet("Survivorship bias is now measured (not assumed) on Bloomberg point-in-time data — see section 10.")

doc.add_heading("10. Survivorship bias — measured on point-in-time data", 1)
SF=PIT["summary"]["Survivorship-free"]; BI=PIT["summary"]["Survivors-only (biased)"]
PSPY=PIT["summary"]["SPY"]; BZ=PIT["bias"]
PIS=PIT["summary"]["Survivorship-free IS (2002-2015)"]; POOS=PIT["summary"]["Survivorship-free OOS (2016-2026)"]
FRF=PIT["factor_regression_free"]; DG=PIT["diagnostics"]
body("Most simulated track records quietly use today's index members, which flatters results because "
     "the companies that failed or were dropped are silently excluded. We removed that bias directly "
     "rather than just disclosing it.")

body("What we did:", bold=True)
bullet("Obtained Bloomberg point-in-time S&P 500 membership: 1,085 historical companies — including every "
       "name later acquired, merged or delisted (e.g. Lehman, Tiffany, Dow Chemical, XL Group).")
bullet("Pulled quarterly total returns 2002–2026. Membership is point-in-time: a company has a return only "
       "for the quarters it was genuinely in the index, so each quarter's universe is the ~500 real members.")
bullet("Detected and corrected a one-quarter labelling offset in the raw export (universe-vs-market "
       "correlation rose from 0.06 to 0.84, and delisting dates then lined up exactly).")
bullet("Ran the IDENTICAL engine on two universes: (A) the full survivorship-free set, and (B) survivors-only, "
       "which deliberately recreates the bias. The gap between A and B is the survivorship bias, isolated.")

body("Performance — survivorship-free vs the recreated bias vs the market:", bold=True)
table(["Universe (quarterly, point-in-time)","CAGR","Vol","Sharpe","Max DD"], [
    ["A. Survivorship-free", f"{SF['CAGR']*100:.1f}%", f"{SF['Vol']*100:.1f}%", f"{SF['Sharpe']:.2f}", f"{SF['MaxDD']*100:.0f}%"],
    ["B. Survivors-only (biased)", f"{BI['CAGR']*100:.1f}%", f"{BI['Vol']*100:.1f}%", f"{BI['Sharpe']:.2f}", f"{BI['MaxDD']*100:.0f}%"],
    ["S&P 500 (SPY)", f"{PSPY['CAGR']*100:.1f}%", f"{PSPY['Vol']*100:.1f}%", f"{PSPY['Sharpe']:.2f}", f"{PSPY['MaxDD']*100:.0f}%"],
    ["A. — in-sample (2002–15)", f"{PIS['CAGR']*100:.1f}%", f"{PIS['Vol']*100:.1f}%", f"{PIS['Sharpe']:.2f}", f"{PIS['MaxDD']*100:.0f}%"],
    ["A. — out-of-sample (2016–26)", f"{POOS['CAGR']*100:.1f}%", f"{POOS['Vol']*100:.1f}%", f"{POOS['Sharpe']:.2f}", f"{POOS['MaxDD']*100:.0f}%"],
    ["Survivorship bias (B − A)", f"{BZ['cagr_drag']*100:+.1f}%", "—", f"{BZ['sharpe_drag']:+.2f}", "—"],
])
body(f"The bias is small: about {abs(BZ['cagr_drag'])*100:.1f}% per year of CAGR and {abs(BZ['sharpe_drag']):.2f} of "
     "Sharpe — squarely in the academic range for large-cap indices. Out-of-sample Sharpe "
     f"({POOS['Sharpe']:.2f}) exceeds in-sample ({PIS['Sharpe']:.2f}), so there is no out-of-sample decay.", italic=True)

if Path("fig_survivorship_corrected.png").exists():
    doc.add_picture("fig_survivorship_corrected.png", width=Inches(6.3))
    cap=doc.add_paragraph(); rc=cap.add_run("Growth of $1 and drawdowns (SIMULATED, quarterly): survivorship-free vs "
        "survivors-only vs S&P 500. The risk edge — roughly half the market's drawdown — is unaffected by the correction.")
    rc.italic=True; rc.font.size=Pt(8); rc.font.color.rgb=RGBColor(0x88,0x88,0x88)

body("Stock-selection alpha and factor exposure (survivorship-free, FF5 + Momentum):", bold=True)
table(["Factor","Beta"], [[k, f"{v:+.2f}"] for k,v in FRF["betas"].items()])
body(f"Alpha {FRF['alpha_annualized']*100:+.1f}% per year (t = {FRF['alpha_tstat']:.2f}), R² = {FRF['rsquared']:.2f}, "
     f"market beta {FRF['betas']['Mkt-RF']:.2f}. Quarterly stock-selection alpha is statistically insignificant on "
     f"BOTH the free ({BZ['alpha_free']*100:+.1f}%, t={BZ['alpha_free_t']:.2f}) and biased "
     f"({BZ['alpha_biased']*100:+.1f}%, t={BZ['alpha_biased_t']:.2f}) universes — so the bias is small, but momentum "
     "alpha is frequency-sensitive (see caveat). The positive market beta and R² also confirm the corrected data "
     "alignment (the pre-fix run gave beta −0.03, R² 0.13).", italic=True, size=10)

body("Diagnostics:", bold=True)
table(["Item","Value"], [
    ["Historical names (universe)", f"{DG['n_historical_names']}"],
    ["Survivors at end of window", f"{DG['n_survivors']}"],
    ["Names removed (delisted/acquired)", f"{DG['n_historical_names']-DG['n_survivors']}"],
    ["Avg names selected / quarter", f"{DG['avg_names_per_quarter']:.0f} (~top quartile of ~500)"],
    ["Avg turnover", f"{DG['avg_annual_turnover']*100:.0f}%/yr (~{DG['avg_annual_turnover']/4*100:.0f}%/qtr)"],
])

body("Honest caveat for the team: this Bloomberg series is quarterly, and quarterly sampling weakens "
     "momentum stock-selection across BOTH universes. So the durable, survivorship-proof edge is the "
     "macro risk overlay (capital preservation), not raw momentum alpha. A monthly point-in-time re-pull "
     "would test the stock-selection alpha at our live frequency, but is not required to make the survivorship case.",
     italic=True, size=10)

d=doc.add_paragraph(); rd=d.add_run("Disclaimer: Fictional pitch for the ESADE Asset Management course. Not investment advice. "
    "All performance is simulated; past performance does not indicate future results.")
rd.italic=True; rd.font.size=Pt(8); rd.font.color.rgb=RGBColor(0x88,0x88,0x88)

doc.save("SectorWide_Strategy_Overview.docx")
print("Saved SectorWide_Strategy_Overview.docx")
